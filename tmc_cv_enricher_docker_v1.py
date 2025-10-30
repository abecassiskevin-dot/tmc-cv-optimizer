#!/usr/bin/env python3
"""
TMC Universal CV Enricher
Lit n'importe quel CV â†’ Enrichit avec IA â†’ GÃ©nÃ¨re CV TMC professionnel
"""

import os
import sys
import json
from docxtpl import DocxTemplate, RichText
from docx import Document
import jinja2
from typing import Dict, List, Any
import PyPDF2
import re
from zipfile import ZipFile
from xml.etree import ElementTree as ET

# === NOUVEAUX IMPORTS POUR OCR ===
from pdf2image import convert_from_path
import pytesseract
from PIL import Image
import tempfile

print(">>> tmc_universal_enricher module loading", flush=True)


def fix_table_width_to_auto(doc):
    """
    Change table width from fixed to auto to prevent horizontal shift after merge.
    
    This fixes the issue where Skills Matrix tables with fixed width (e.g., 8.1 inches)
    get shifted right after merging because they don't fit within the page margins.
    
    Args:
        doc: Document object to fix
    
    Returns:
        int: Number of tables fixed
    """
    w = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    tables_fixed = 0
    
    for table in doc.tables:
        tbl = table._element
        tblPr = tbl.find(f'.//{w}tblPr')
        
        if tblPr is not None:
            # Find and fix tblW (table width)
            tblW = tblPr.find(f'.//{w}tblW')
            if tblW is not None:
                old_type = tblW.get(f'{w}type', 'unknown')
                old_w = tblW.get(f'{w}w', 'unknown')
                
                # Change to auto width
                tblW.set(f'{w}type', 'auto')
                tblW.set(f'{w}w', '0')
                
                print(f"   ğŸ”§ Table width changed: {old_type}={old_w} â†’ auto=0")
                tables_fixed += 1
            
            # Remove fixed layout if present
            tblLayout = tblPr.find(f'.//{w}tblLayout')
            if tblLayout is not None:
                old_layout = tblLayout.get(f'{w}type', 'unknown')
                tblPr.remove(tblLayout)
                print(f"   ğŸ”§ Removed tblLayout: {old_layout}")
    
    return tables_fixed


class TMCUniversalEnricher:
    """Enrichisseur universel de CV au format TMC"""
    
    def __init__(self, api_key: str = None):
        """Initialiser avec clÃ© API Claude"""
        self.api_key = api_key or os.getenv('ANTHROPIC_API_KEY')
        if not self.api_key:
            raise ValueError("âŒ ClÃ© API Claude manquante! DÃ©finissez ANTHROPIC_API_KEY dans les secrets Streamlit ou en variable d'environnement.")
        
        # Debug clÃ© API
        print(f">>> ANTHROPIC_KEY_PRESENT: {bool(self.api_key)}, len: {len(self.api_key) if self.api_key else 0}", flush=True)
        
        # Ne crÃ©e PAS le client ici (lazy loading)
        self._anthropic_client = None
    
    def _get_anthropic_client(self):
        """Lazy loading du client Anthropic"""
        if self._anthropic_client is None:
            try:
                print(">>> Creating anthropic client", flush=True)
                import anthropic
                # CrÃ©ation SIMPLE du client pour version 0.25.9
                self._anthropic_client = anthropic.Anthropic(api_key=self.api_key)
                print(">>> Anthropic client created OK", flush=True)
            except Exception as e:
                print(f">>> ERROR creating anthropic client: {repr(e)}", flush=True)
                raise
        return self._anthropic_client
    
    # ========================================
    # MODULE 1 : EXTRACTION UNIVERSELLE
    # ========================================
    
    def detect_file_type(self, file_path: str) -> str:
        """DÃ©tecter le type de fichier"""
        ext = file_path.lower().split('.')[-1]
        if ext == 'pdf':
            return 'pdf'
        elif ext in ['docx', 'doc']:
            return 'docx'
        elif ext in ['txt', 'text']:
            return 'txt'
        else:
            return 'unknown'
    
    def extract_from_pdf(self, file_path: str) -> str:
        """
        Extraire texte d'un PDF avec fallback OCR automatique
        1. Essaye PyPDF2 pour texte sÃ©lectionnable
        2. Si Ã©chec/texte vide â†’ Utilise OCR sur images
        """
        print(f"ğŸ“„ Extracting PDF: {file_path}", flush=True)
        
        try:
            # ===== Ã‰TAPE 1: Tentative extraction PyPDF2 =====
            text = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                print(f"ğŸ“Š PDF has {num_pages} pages", flush=True)
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
                    print(f"  Page {page_num}: {len(page_text) if page_text else 0} chars", flush=True)
            
            extracted_text = "\n".join(text).strip()
            
            # ===== VÃ‰RIFIER SI L'EXTRACTION A FONCTIONNÃ‰ =====
            # Seuil: Si moins de 100 caractÃ¨res ou trop peu de mots â†’ C'est scannÃ©
            word_count = len(extracted_text.split())
            char_count = len(extracted_text)
            
            print(f"ğŸ“ˆ PyPDF2 extraction: {char_count} chars, {word_count} words", flush=True)
            
            # Si extraction suffisante â†’ Retourner
            if char_count > 100 and word_count > 20:
                print("âœ… PDF text extraction successful (text-based PDF)", flush=True)
                return extracted_text
            
            # ===== Ã‰TAPE 2: PDF scannÃ© dÃ©tectÃ© â†’ OCR =====
            print("âš ï¸ PDF appears to be scanned (image-based). Switching to OCR...", flush=True)
            return self._extract_from_pdf_ocr(file_path)
            
        except Exception as e:
            print(f"âŒ Error in PDF extraction: {e}", flush=True)
            # En cas d'erreur PyPDF2, essayer quand mÃªme OCR
            try:
                print("ğŸ”„ Trying OCR as fallback...", flush=True)
                return self._extract_from_pdf_ocr(file_path)
            except Exception as e2:
                print(f"âŒ OCR fallback also failed: {e2}", flush=True)
                return ""
    
    def _extract_from_pdf_ocr(self, file_path: str) -> str:
        """
        Extraire texte d'un PDF scannÃ© via OCR
        Utilise pdf2image + pytesseract
        """
        print("ğŸ” Starting OCR extraction...", flush=True)
        
        try:
            # Convertir PDF en images (une par page)
            # poppler_path peut Ãªtre nÃ©cessaire sur Windows, mais pas sur Linux/Render
            images = convert_from_path(
                file_path,
                dpi=300,  # Haute rÃ©solution pour meilleur OCR
                fmt='jpeg',
                thread_count=2  # ParallÃ©lisation
            )
            
            print(f"ğŸ“· Converted {len(images)} pages to images", flush=True)
            
            # Extraire texte de chaque image
            all_text = []
            for i, image in enumerate(images, 1):
                print(f"  ğŸ” OCR processing page {i}/{len(images)}...", flush=True)
                
                # Appliquer OCR avec config optimisÃ©e
                # lang='eng+fra' pour anglais ET franÃ§ais
                page_text = pytesseract.image_to_string(
                    image,
                    lang='eng+fra',  # Anglais + FranÃ§ais
                    config='--psm 1 --oem 3'  # PSM 1 = automatic page segmentation with OSD
                )
                
                if page_text.strip():
                    all_text.append(f"--- Page {i} ---\n{page_text}")
                    print(f"  âœ“ Page {i}: {len(page_text)} chars extracted", flush=True)
            
            extracted_text = "\n\n".join(all_text)
            print(f"âœ… OCR extraction complete: {len(extracted_text)} chars total", flush=True)
            
            return extracted_text
            
        except Exception as e:
            print(f"âŒ OCR extraction failed: {e}", flush=True)
            import traceback
            print(traceback.format_exc(), flush=True)
            return ""
    
     
    def extract_from_docx(self, file_path: str) -> str:
        """Extraire texte d'un Word + zones textes"""
        try:
            doc = Document(file_path)
            text = []
            
            # Paragraphes normaux
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text.strip())
            
            # Tableaux
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text.append(row_text)
            
            # NOUVEAU : Extraire les zones textes du XML
            textbox_content = self.extract_textboxes(file_path)
            if textbox_content:
                text.append("\n=== ZONES TEXTES ===")
                text.extend(textbox_content)
            
            return "\n".join(text)
        except Exception as e:
            print(f"âš ï¸ Erreur extraction Word: {e}")
            return ""
    def extract_from_txt(self, file_path: str) -> str:
        """Extraire texte d'un fichier texte"""
        try:
            # Essayer plusieurs encodages
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            # Si tout Ã©choue, ignorer les erreurs
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            print(f"âš ï¸ Erreur extraction TXT: {e}")
            return ""
    
    def extract_textboxes(self, docx_path: str) -> list:
        """Extraire le contenu des zones textes (text boxes) du XML"""
        textboxes = []
        
        try:
            with ZipFile(docx_path, 'r') as docx:
                # Lire le document.xml
                xml_content = docx.read('word/document.xml')
                tree = ET.fromstring(xml_content)
                
                # Namespaces Word
                namespaces = {
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                    'v': 'urn:schemas-microsoft-com:vml',
                    'w10': 'urn:schemas-microsoft-com:office:word'
                }
                
                # Chercher tous les Ã©lÃ©ments de texte dans les zones textes
                # Les zones textes sont dans w:txbxContent
                for txbx in tree.findall('.//w:txbxContent', namespaces):
                    texts = []
                    for t in txbx.findall('.//w:t', namespaces):
                        if t.text:
                            texts.append(t.text.strip())
                    if texts:
                        textboxes.append(' '.join(texts))
                
                # Aussi chercher dans v:textbox (ancien format)
                for vtxbx in tree.findall('.//v:textbox', namespaces):
                    texts = []
                    for t in vtxbx.findall('.//w:t', namespaces):
                        if t.text:
                            texts.append(t.text.strip())
                    if texts:
                        textboxes.append(' '.join(texts))
                        
        except Exception as e:
            print(f"âš ï¸ Erreur extraction zones textes: {e}")
        
        return textboxes
    
    def extract_cv_text(self, cv_path: str) -> str:
        """Extraction universelle - dÃ©tecte et extrait selon le type"""
        print(f"ğŸ“„ Extraction du CV: {cv_path}")
        
        file_type = self.detect_file_type(cv_path)
        
        if file_type == 'pdf':
            print("   Format dÃ©tectÃ©: PDF")
            return self.extract_from_pdf(cv_path)
        elif file_type == 'docx':
            print("   Format dÃ©tectÃ©: Word")
            return self.extract_from_docx(cv_path)
        elif file_type == 'txt':
            print("   Format dÃ©tectÃ©: Texte")
            return self.extract_from_txt(cv_path)
        else:
            raise ValueError(f"âŒ Format non supportÃ©: {file_type}")

    # ========================================
    # MODULE 2 : PARSING INTELLIGENT
    # ========================================
    
    def parse_cv_with_claude(self, cv_text: str) -> Dict[str, Any]:
        """Parser le CV avec Claude pour extraire les infos structurÃ©es"""
        print("ğŸ¤– Parsing du CV avec Claude AI...", flush=True)
        
        try:
            client = self._get_anthropic_client()
            
            prompt = f"""Tu es un expert en analyse de CV. Extrait TOUTES les informations de ce CV et structure-les en JSON.

CV Ã€ ANALYSER:
{cv_text}

IMPORTANT CRITIQUE:
- Le NOM peut Ãªtre cachÃ© dans un tableau HTML ou Ãªtre stylisÃ©. Cherche PARTOUT.
- Le LIEU DE RÃ‰SIDENCE est OBLIGATOIRE : cherche "MontrÃ©al", "Montreal", villes + pays (ex: "Montreal CA", "MontrÃ©al, Canada", "Toronto ON", etc.). Si introuvable, mets "Location not specified".
- Les LANGUES sont OBLIGATOIRES : cherche "FranÃ§ais", "French", "English", "Anglais", "Bilingual", "Bilingue", etc. Si introuvable, mets ["Not specified"].

Extrait et structure en JSON STRICT (sans markdown):
{{
  "nom_complet": "Nom PrÃ©nom du candidat (cherche PARTOUT, mÃªme dans tableaux/HTML)",
  "titre_professionnel": "Titre/poste actuel",
  "profil_resume": "RÃ©sumÃ© du profil si prÃ©sent (sinon vide)",
  "lieu_residence": "OBLIGATOIRE - Ville, Pays (ex: MontrÃ©al, Canada) ou Montreal CA. Cherche codes pays (CA, US, FR). Si vraiment introuvable: 'Location not specified'",
  "langues": ["OBLIGATOIRE - FranÃ§ais", "Anglais", ... Cherche 'bilingual', 'French', 'English', etc. Si introuvable: ['Not specified']],
  "competences": ["compÃ©tence1", "compÃ©tence2", "compÃ©tence3", ...],
  "experiences": [
    {{
      "periode": "2020-2023",
      "entreprise": "Nom entreprise",
      "poste": "Titre du poste",
      "responsabilites": ["tÃ¢che 1", "tÃ¢che 2", "tÃ¢che 3"]
    }}
  ],
  "formation": [
    {{
      "diplome": "Nom COMPLET du diplÃ´me",
      "institution": "Nom Ã©cole/universitÃ©",
      "annee": "2020 (ou pÃ©riode exacte)",
      "pays": "Canada"
    }}
  ],
  "certifications": [
    {{
      "nom": "Nom certification",
      "organisme": "Organisme",
      "annee": "2023"
    }}
  ],
  "projets": [
    {{
      "nom": "Nom projet",
      "description": "Description courte"
    }}
  ]
}}

RÃˆGLES CRITIQUES:
- Le NOM est PRIORITAIRE - cherche dans tout le texte (tableaux, dÃ©but, fin)
- LIEU DE RÃ‰SIDENCE : cherche formats "Ville, Pays", "Montreal CA", "MontrÃ©al QC", codes postaux (H2X, etc.)
- LANGUES : cherche "Languages", "Langues", "French", "English", "Bilingual", mÃªme dans sections compÃ©tences
- Pour les diplÃ´mes: nom COMPLET + annÃ©e EXACTE
- Extrait TOUT (ne rate rien)
- Si une section est vide, mets une liste vide []
- Format JSON strict uniquement"""

            print(f">>> Calling Claude API with timeout=300s...", flush=True)
            response = client.messages.create(
                model="claude-sonnet-4-5-20250929",
                max_tokens=8000,
                timeout=300.0,  # 5 minutes max
                messages=[{"role": "user", "content": prompt}]
            )
            print(f">>> API call completed successfully", flush=True)
            
        except Exception as e:
            print(f">>> ERROR calling anthropic for parsing: {repr(e)}", flush=True)
            return {}
        
        response_text = response.content[0].text.strip()
        
        # Nettoyer JSON
        if response_text.startswith('```json'):
            response_text = response_text[7:]
        if response_text.startswith('```'):
            response_text = response_text[3:]
        if response_text.endswith('```'):
            response_text = response_text[:-3]
        response_text = response_text.strip()
        
        try:
            parsed_data = json.loads(response_text)
            print(f"âœ… Parsing rÃ©ussi!")
            print(f"   Nom: [ANONYMIZED]")
            print(f"   Langues: {', '.join(parsed_data.get('langues', []))}")
            print(f"   Lieu: [ANONYMIZED]")
            print(f"   CompÃ©tences: {len(parsed_data.get('competences', []))}")
            print(f"   ExpÃ©riences: {len(parsed_data.get('experiences', []))}")
            return parsed_data
        except json.JSONDecodeError as e:
            print(f"âš ï¸ Erreur JSON: {e}")
            print(f"RÃ©ponse brute: {response_text[:500]}")
            return {}

    # ========================================
    # MODULE 3 : ENRICHISSEMENT (TON PROMPT)
    # ========================================
    
    def read_job_description(self, jd_path: str) -> str:
        """Lire la job description"""
        file_type = self.detect_file_type(jd_path)
        
        if file_type == 'pdf':
            return self.extract_from_pdf(jd_path)
        elif file_type == 'docx':
            return self.extract_from_docx(jd_path)
        else:
            return self.extract_from_txt(jd_path)
    
    def analyze_cv_matching(self, parsed_cv: Dict[str, Any], jd_text: str) -> Dict[str, Any]:
        """
        Analyser le matching entre CV et JD sans enrichir le contenu.
        Retourne uniquement: score_matching, domaines_analyses, synthese_matching
        """
        import time
        
        print(f"ğŸ” Analyse du matching CV/JD...", flush=True)
        
        start_time = time.time()
        
        try:
            client = self._get_anthropic_client()
            
            # Reconstruire le CV en texte pour le prompt
            cv_text = f"""
PROFIL: {parsed_cv.get('profil_resume', '')}

TITRE: {parsed_cv.get('titre_professionnel', '')}

COMPÃ‰TENCES:
{chr(10).join(['- ' + comp for comp in parsed_cv.get('competences', [])])}

EXPÃ‰RIENCES:
"""
            for exp in parsed_cv.get('experiences', []):
                cv_text += f"\n{exp.get('periode', '')} | {exp.get('entreprise', '')} | {exp.get('poste', '')}\n"
                for resp in exp.get('responsabilites', []):
                    cv_text += f"  - {resp}\n"
            
            cv_text += "\nFORMATION:\n"
            for form in parsed_cv.get('formation', []):
                cv_text += f"- {form.get('diplome', '')} | {form.get('institution', '')} | {form.get('annee', '')}\n"
        
            # PROMPT FOCALISÃ‰ SUR L'ANALYSE DE MATCHING UNIQUEMENT - VERSION ULTRA-STRICTE V1.3.9
            prompt = f"""Tu es un systÃ¨me d'Ã©valuation automatisÃ© ULTRA-STRICT qui analyse le matching entre CV et Job Description.

ğŸ¯ ANALYSE DE MATCHING PONDÃ‰RÃ‰E (VERSION ULTRA-STRICTE V1.3.9):

âš ï¸ PRINCIPE FONDAMENTAL - Ã‰VALUATION ULTRA-RIGOUREUSE:
- Tu es un RECRUTEUR SENIOR EXTRÃŠMEMENT EXIGEANT avec 15+ ans d'expÃ©rience
- Tu recrutes pour des postes CRITIQUES oÃ¹ l'excellence est la norme
- CHAQUE point doit Ãªtre MÃ‰RITÃ‰ avec des PREUVES CONCRÃˆTES du CV
- Si tu hÃ©sites entre 2 scores â†’ TOUJOURS prends le PLUS BAS
- Agis comme si tu recrutais pour ton propre argent (zÃ©ro tolÃ©rance pour l'approximation)
- Pour le MÃŠME CV et la MÃŠME JD â†’ EXACTEMENT le mÃªme score Ã  chaque fois (cohÃ©rence algorithmique)

ğŸ”´ RÃˆGLE D'OR - SCORE GLOBAL = SOMME DOMAINES:
- Le score_matching FINAL = somme EXACTE de tous les scores de domaines
- VÃ‰RIFIE 3 FOIS avant de rÃ©pondre: somme des scores = score_matching
- Si tu calcules 58/100 en sommant les domaines â†’ score_matching DOIT Ãªtre 58
- NE JAMAIS inventer un score global diffÃ©rent de la somme calculÃ©e

â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
ğŸ“‹ Ã‰TAPE 1 - IDENTIFIER 5-8 DOMAINES CRITIQUES
â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

PROCESSUS AUTOMATIQUE D'IDENTIFICATION:
1. Scan complet de la JD - repÃ©rer TOUS les mots techniques/compÃ©tences
2. Compter la frÃ©quence EXACTE de chaque technologie/compÃ©tence/mÃ©thodologie
3. Identifier les must-haves vs nice-to-haves
4. CrÃ©er une liste de domaines par ordre d'importance
5. Appliquer la formule de pondÃ©ration ci-dessous

ğŸ“Š FORMULE DE PONDÃ‰RATION MATHÃ‰MATIQUE:
Pour chaque domaine, calcule son poids avec:
Poids = (Mentions_JD Ã— 10) + (Niveau_requis Ã— 5) + Bonus_contexte

OÃ¹:
- Mentions_JD: Nombre de fois mentionnÃ© dans JD (1=once, 2=2-3 times, 3=4+ times)
- Niveau_requis: Must-have/Required=3, Important=2, Nice-to-have=1
- Bonus_contexte: +5 si dans le titre du poste, +3 si dans top requirements

ğŸ’¡ EXEMPLES DE DOMAINES TYPES:
- Technologies spÃ©cifiques (ex: "Python Django", "AWS Lambda", "React Native")
- MÃ©thodologies (ex: "Agile/Scrum", "ITIL v4", "DevOps CI/CD")
- CompÃ©tences mÃ©tier (ex: "Financial Modeling", "Clinical Trials Management")
- Certifications (ex: "PMP", "AWS Solutions Architect", "CPA")
- Langues avec niveau (ex: "Bilingual French/English C1+", "Spanish Business Level")
- Soft skills MESURABLES (ex: "Team Leadership 10+ people", "Stakeholder Management C-Suite")

âš ï¸ INTERDICTIONS ABSOLUES:
- NE JAMAIS crÃ©er de domaine vague type "General Fit", "Soft Skills", "Cultural Fit"
- NE JAMAIS crÃ©er de domaine "bonus" pour ajuster artificiellement le score
- TOUS les domaines doivent Ãªtre EXPLICITEMENT mentionnÃ©s dans la JD
- Pas de domaines "catch-all" ou gÃ©nÃ©riques

â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
ğŸ¯ Ã‰TAPE 2 - GRILLE D'Ã‰VALUATION ULTRA-STRICTE
â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

Pour CHAQUE domaine identifiÃ©, Ã©value le score avec cette GRILLE ULTRA-SÃ‰VÃˆRE (0-100 points par domaine):

â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”
ğŸ”´ NIVEAU 0-15 POINTS: QUASI-AUCUNE COMPÃ‰TENCE
â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”
0 points: CompÃ©tence TOTALEMENT absente du CV (aucune mention directe ou indirecte)
10 points: Mention trÃ¨s vague OU compÃ©tence tangentielle (ex: "exposure to", "familiar with")
15 points: Mention superficielle OU formation thÃ©orique seulement SANS pratique OU <3 mois d'expÃ©rience

ğŸŸ  NIVEAU 20-35 POINTS: DÃ‰BUTANT/JUNIOR
â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”
20 points: 3-6 mois d'expÃ©rience pratique OU 1 projet simple rÃ©alisÃ© sous supervision
25 points: 6-9 mois d'expÃ©rience OU 2 projets avec support d'Ã©quipe
30 points: 9-12 mois d'expÃ©rience avec autonomie partielle OU certification rÃ©cente + pratique limitÃ©e
35 points: 1 an d'expÃ©rience solide avec quelques rÃ©alisations concrÃ¨tes (mais sans metrics)

ğŸŸ¡ NIVEAU 40-55 POINTS: INTERMÃ‰DIAIRE
â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”
40 points: 1-1.5 ans d'expÃ©rience + 2-3 projets pertinents documentÃ©s
45 points: 1.5-2 ans d'expÃ©rience + contribution mesurable (ex: "improved X by Y%")
50 points: 2-2.5 ans d'expÃ©rience solide + rÃ©alisations quantifiÃ©es (metrics, budget, scope)
55 points: 2.5-3 ans + rÃ´le de contributeur principal sur projets moyens

ğŸŸ¢ NIVEAU 60-75 POINTS: CONFIRMÃ‰/SENIOR
â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”
60 points: 3-4 ans d'expÃ©rience confirmÃ©e + ownership de projets + rÃ©sultats mesurables
65 points: 4-5 ans + expertise dÃ©montrÃ©e par rÃ©alisations significatives (ex: led team of 5, managed $500K budget)
70 points: 5-6 ans + rÃ´le de lead/expert technique + mentorship + process improvements
75 points: 6-7 ans + expertise reconnue EN INTERNE (promotions, leadership technique, formations donnÃ©es en interne)

ğŸ”µ NIVEAU 80-90 POINTS: EXPERT EXCEPTIONNEL
â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”
80 points: 7-8 ans d'expÃ©rience TRÃˆS solide + leadership prouvÃ© + expertise reconnue PAR L'INDUSTRIE (speaking engagements, certifications avancÃ©es, articles techniques)
85 points: 8-10 ans + contribution MAJEURE Ã  l'industrie (architecture de solutions complexes multi-millions, thought leadership, certifications rares)
90 points: 10-12 ans + expertise de NIVEAU MONDIAL dans ce domaine spÃ©cifique (publications acadÃ©miques/industrie, confÃ©rences internationales, mentor d'experts, awards/recognition)

ğŸ† NIVEAU 95-100 POINTS: QUASI-IMPOSSIBLE - TOP 0.1% MONDIAL
â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”
95 points: 12-15 ans + reconnaissance INTERNATIONALE + contributions MAJEURES Ã  l'Ã©volution du domaine (patents, standards, books, keynote speaker top conferences)
100 points: RÃ‰SERVÃ‰ AUX LÃ‰GENDES VIVANTES - 15+ ans + autoritÃ© MONDIALE incontestÃ©e dans le domaine + impact transformationnel sur l'industrie (ex: crÃ©ateur de framework utilisÃ© par millions, membre de comitÃ©s internationaux, consultant pour Fortune 10)

âš ï¸ RÃˆGLES ULTRA-STRICTES D'ATTRIBUTION:
â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”
1. JAMAIS de score â‰¥60 sans PREUVES QUANTIFIÃ‰ES concrÃ¨tes dans le CV
2. JAMAIS de score â‰¥75 sans leadership/mentorship/expertise reconnue PROUVÃ‰E
3. JAMAIS de score â‰¥85 sans contributions MAJEURES Ã  l'industrie (publications, speaking, thought leadership)
4. JAMAIS de score â‰¥95 sans reconnaissance INTERNATIONALE vÃ©rifiable
5. Si le CV mentionne l'expÃ©rience en annÃ©es SEULEMENT sans dÃ©tails de rÃ©alisations â†’ score MAX = 55
6. Si aucun metric/chiffre fourni pour un domaine â†’ score MAX = 50
7. Si le candidat change de domaine/technologie frÃ©quemment (job hopping) â†’ pÃ©nalitÃ© de -10 points
8. Certifications SANS expÃ©rience pratique associÃ©e â†’ score MAX = 30
9. ExpÃ©rience dans environnement non-professionnel (side projects, freelance) compte pour 50% seulement
10. Si tu hÃ©sites entre 2 scores â†’ TOUJOURS choisir le PLUS BAS

âš™ï¸ RÃˆGLES DE CALCUL FINAL:
1. Score brut du domaine = Ã©valuation selon grille ci-dessus (0-100)
2. Score pondÃ©rÃ© = (score_brut Ã— poids) / 100
3. Score_max du domaine = poids

Exemple dÃ©taillÃ©:
- Domaine: "Python Backend Development" | Poids: 25%
- Candidat: 4.5 ans d'expÃ©rience Python, 3 projets documentÃ©s, led team of 3, aucune publication
- Ã‰valuation: Entre 60 et 65 points â†’ choisir 60 (rÃ¨gle du plus bas)
- Score pondÃ©rÃ©: (60 Ã— 25) / 100 = 15 points
- Score_max: 25 points
- Notation: 15/25

â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
ğŸ“Š Ã‰TAPE 3 - CALCULER LE SCORE TOTAL
â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

Score_matching = SOMME de tous les scores pondÃ©rÃ©s (arrondi Ã  l'entier)

Exemple:
15 (Python) + 10 (AWS) + 8 (Agile) + 12 (API Design) + 9 (PostgreSQL) + 7 (Docker) = 61/100

âš ï¸ VÃ‰RIFICATIONS FINALES OBLIGATOIRES:
â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”
1. Somme des poids = EXACTEMENT 100%
2. Score_matching = somme EXACTE des scores pondÃ©rÃ©s
3. Si score > 80 â†’ TRIPLE-CHECK: y a-t-il vraiment des preuves d'expertise exceptionnelle?
4. Si score > 90 â†’ QUADRUPLE-CHECK: est-ce vraiment un candidat top 1% mondial? (la rÃ©ponse devrait presque toujours Ãªtre NON)
5. Refaire le calcul 2 fois pour confirmer

ğŸ¯ PHILOSOPHIE DE NOTATION ATTENDUE (distribution rÃ©aliste):
â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”
- Score 95-100: <1% des candidats (quasi-impossible, rÃ©servÃ© aux lÃ©gendes)
- Score 85-94: ~5% (top performers exceptionnels)
- Score 75-84: ~15% (trÃ¨s bons candidats confirmÃ©s)
- Score 65-74: ~25% (bons candidats solides)
- Score 50-64: ~30% (candidats acceptables avec gaps)
- Score <50: ~24% (candidats insuffisants)

âš ï¸ DERNIÃˆRE VÃ‰RIFICATION AVANT RÃ‰PONSE:
â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”
Pose-toi ces questions pour CHAQUE domaine oÃ¹ tu as donnÃ© â‰¥60 points:
- Ai-je des PREUVES CONCRÃˆTES d'expÃ©rience quantifiable dans le CV?
- Ai-je des RÃ‰ALISATIONS MESURABLES (metrics, budget, team size, impact)?
- Le candidat a-t-il eu un rÃ´le de LEADERSHIP/OWNERSHIP dÃ©montrÃ©?
- Pour les scores â‰¥85: y a-t-il des contributions Ã  l'INDUSTRIE (publications, speaking, thought leadership)?
Si la rÃ©ponse n'est pas un OUI catÃ©gorique avec preuves multiples â†’ BAISSE le score.

â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
ğŸ“ Ã‰TAPE 4 - SYNTHÃˆSE EXECUTIVE (4-5 LIGNES MAX)
â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

RÃ©dige une synthÃ¨se ULTRA-CONCISE en 4-5 LIGNES (80-100 mots maximum) qui:

STRUCTURE OBLIGATOIRE (1 paragraphe fluide):
1. Lead with match level + score (e.g., "GOOD match (73/100) for [Role]")
2. Highlight 2-3 TOP strengths with brief evidence (years, key achievement, metric)
3. Mention 1-2 minor gaps or "nice-to-haves" missing
4. End with clear recommendation: "Interview - [reason]" or "Pass - [reason]"

EXEMPLE FORMAT:
"GOOD match (73/100) for Senior Full-Stack Developer. Strong Python backend (8 years) with proven cloud migration leadership (60% deployment time reduction). Full-stack capability confirmed with React + modern DevOps. Minor gaps: Kubernetes nice-to-have, limited Montreal-specific experience. Recommendation: Interview - solid technical fit with measurable impact."

RÃˆGLES CRITIQUES:
- MAX 4-5 lignes (80-100 mots)
- NO paragraphs, NO bullet points - juste 1 bloc de texte fluide
- Include score + match level (EXCELLENT 85+, GOOD 70-84, MODERATE 55-69, WEAK <55)
- Be specific with numbers/metrics when available
- Professional but direct tone
- Clear go/no-go recommendation at the end

â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
ğŸ“„ FORMAT DE SORTIE JSON
â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

ğŸ“„ JOB DESCRIPTION:
{jd_text}

ğŸ“„ CV DU CANDIDAT:
{cv_text}

â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

ğŸ¯ GÃ‰NÃˆRE MAINTENANT TON ANALYSE - FORMAT JSON STRICT:

Retourne UNIQUEMENT un JSON avec cette structure (sans texte avant/aprÃ¨s):

{{
    "score_matching": 58,
    "domaines_analyses": [
        {{
            "domaine": "Nom du domaine technique/compÃ©tence exact",
            "poids": 20,
            "score": 10,
            "score_max": 20,
            "match": "bon",
            "commentaire": "Justification FACTUELLE ultra-dÃ©taillÃ©e basÃ©e sur des Ã©lÃ©ments PRÃ‰CIS du CV avec annÃ©es d'expÃ©rience, projets, rÃ©alisations, metrics. Minimum 2-3 phrases complÃ¨tes."
        }}
    ],
    "synthese_matching": "COMPREHENSIVE PROFESSIONAL ANALYSIS (4-6 DETAILED PARAGRAPHS, 250-350 WORDS):

[Paragraph 1 - Overall Assessment]
[Detailed assessment text...]

[Paragraph 2 - Top Strengths]
[Detailed strengths text...]

[Paragraph 3 - Partial Matches]
[Detailed partial matches text...]

[Paragraph 4 - Gaps]
[Detailed gaps text...]

[Paragraph 5 - Final Recommendation]
[Detailed recommendation text...]"
}}

âš ï¸ RÃˆGLES JSON CRITIQUES:
â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”
- "match" peut Ãªtre: "excellent" (â‰¥85/100), "bon" (65-84), "partiel" (40-64), "incompatible" (<40)
- Tous les scores doivent Ãªtre des NOMBRES (pas de strings)
- La somme des poids doit faire exactement 100
- Le score_matching doit Ãªtre la somme exacte des scores de tous les domaines
- Commentaire: minimum 2-3 phrases complÃ¨tes avec dÃ©tails factuels prÃ©cis du CV
- SynthÃ¨se: MAXIMUM 4-5 lignes (80-100 mots), format executive summary

âš ï¸ LANGUE: ALL output must be in ENGLISH.
- Domain names in English (e.g., "Python Backend Development", not "DÃ©veloppement Backend Python")
- All comments in English
- Synthesis in English (4-5 lines max)

GÃ©nÃ¨re l'analyse maintenant:"""
            
            print(f">>> Calling Claude API for matching analysis...", flush=True)
            
            # âœ… RETRY LOGIC FOR TIMEOUTS
            max_retries = 2
            response = None
            last_error = None
            
            for attempt in range(max_retries):
                try:
                    response = client.messages.create(
                        model="claude-sonnet-4-5-20250929",
                        max_tokens=4000,
                        timeout=900.0,  # 15 minutes
                        messages=[{"role": "user", "content": prompt}]
                    )
                    break  # Success - exit retry loop
                    
                except Exception as e:
                    last_error = e
                    error_name = type(e).__name__
                    
                    # Check if it's a timeout error
                    if 'timeout' in error_name.lower() or 'timeout' in str(e).lower():
                        if attempt < max_retries - 1:
                            print(f"â±ï¸ Timeout attempt {attempt+1}/{max_retries}, retrying...", flush=True)
                            continue
                        else:
                            # Final timeout - return error result
                            print(f"âŒ Final timeout after {max_retries} attempts", flush=True)
                            return {
                                'error': 'timeout',
                                'score_matching': 0,
                                'domaines_analyses': [],
                                'synthese_matching': "â±ï¸ L'analyse a pris trop de temps (timeout aprÃ¨s plusieurs tentatives). Veuillez rÃ©essayer avec un CV plus court ou contactez le support."
                            }
                    else:
                        # Other error - re-raise
                        raise
            
            if response is None:
                # Should not happen, but safety check
                raise last_error
            
            # Extraire tokens
            usage = response.usage
            input_tokens = usage.input_tokens
            output_tokens = usage.output_tokens
            total_tokens = input_tokens + output_tokens
            
            print(f">>> API Response received. Tokens: {total_tokens}", flush=True)
            
            # Parser la rÃ©ponse
            response_text = response.content[0].text.strip()
            
            # Nettoyer le JSON
            if response_text.startswith('```json'):
                response_text = response_text[7:]
            if response_text.startswith('```'):
                response_text = response_text[3:]
            if response_text.endswith('```'):
                response_text = response_text[:-3]
            response_text = response_text.strip()
            
            # Parser le JSON
            try:
                matching_result = json.loads(response_text)
                print(f">>> JSON parsed successfully!", flush=True)
                
                # V1.3.4.1 FIX: Recalculer le score_matching pour garantir cohÃ©rence
                # Somme des scores de tous les domaines
                if 'domaines_analyses' in matching_result and matching_result['domaines_analyses']:
                    calculated_score = sum(d.get('score', 0) for d in matching_result['domaines_analyses'])
                    original_score = matching_result.get('score_matching', 0)
                    
                    # Si diffÃ©rence > 2 points, utiliser le score calculÃ©
                    if abs(calculated_score - original_score) > 2:
                        print(f"âš ï¸ Score mismatch detected: Claude={original_score}, Calculated={calculated_score}")
                        print(f"   Using calculated score for consistency: {calculated_score}/100")
                        matching_result['score_matching'] = round(calculated_score)
                    else:
                        # Petite diffÃ©rence acceptable (arrondis)
                        matching_result['score_matching'] = round(calculated_score)
                    
                    # âœ… V1.3.4.2 FIX: CAP SCORE AT 100 MAXIMUM
                    if matching_result['score_matching'] > 100:
                        print(f"âš ï¸ Score exceeded 100: {matching_result['score_matching']} â†’ Capping at 100")
                        matching_result['score_matching'] = 100
                        
            except json.JSONDecodeError as e:
                print(f"âš ï¸ JSON Error: {e}", flush=True)
                print(f">>> Attempting to fix JSON...", flush=True)
                
                # Tentative de rÃ©paration
                fix_prompt = f"""The following JSON is malformed. Please fix it and return ONLY the corrected JSON without any explanation or markdown:

{response_text}

Return the corrected JSON directly:"""
                
                # âœ… RETRY LOGIC FOR JSON FIX
                fix_response = None
                for fix_attempt in range(2):
                    try:
                        fix_response = client.messages.create(
                            model="claude-sonnet-4-5-20250929",
                            max_tokens=4000,
                            timeout=300.0,  # 5 minutes for fix
                            messages=[{"role": "user", "content": fix_prompt}]
                        )
                        break
                    except Exception as fix_error:
                        if 'timeout' in type(fix_error).__name__.lower() or 'timeout' in str(fix_error).lower():
                            if fix_attempt < 1:
                                print(f"â±ï¸ JSON fix timeout, retrying...", flush=True)
                                continue
                            else:
                                # Can't fix JSON - return error
                                return {
                                    'error': 'json_parse_timeout',
                                    'score_matching': 0,
                                    'domaines_analyses': [],
                                    'synthese_matching': "âŒ Erreur de parsing JSON et timeout lors de la correction. Veuillez rÃ©essayer."
                                }
                        else:
                            raise
                
                if fix_response is None:
                    return {
                        'error': 'json_fix_failed',
                        'score_matching': 0,
                        'domaines_analyses': [],
                        'synthese_matching': "âŒ Impossible de corriger le JSON malformÃ©."
                    }
                
                fixed_text = fix_response.content[0].text.strip()
                if fixed_text.startswith('```json'):
                    fixed_text = fixed_text[7:]
                if fixed_text.startswith('```'):
                    fixed_text = fixed_text[3:]
                if fixed_text.endswith('```'):
                    fixed_text = fixed_text[:-3]
                fixed_text = fixed_text.strip()
                
                matching_result = json.loads(fixed_text)
                print(f">>> JSON successfully fixed and parsed!", flush=True)
            
            # Calculer le temps et coÃ»t
            processing_time = round(time.time() - start_time, 2)
            cost_input = (input_tokens / 1_000_000) * 3.0
            cost_output = (output_tokens / 1_000_000) * 15.0
            total_cost = round(cost_input + cost_output, 4)
            
            # Ajouter les mÃ©tadonnÃ©es
            matching_result['_metadata'] = {
                'processing_time_seconds': processing_time,
                'input_tokens': input_tokens,
                'output_tokens': output_tokens,
                'total_tokens': total_tokens,
                'estimated_cost_usd': total_cost
            }
            
            print(f"âœ… Analyse de matching rÃ©ussie!")
            print(f"   Score matching: {matching_result.get('score_matching', 0)}/100")
            print(f"   Domaines analysÃ©s: {len(matching_result.get('domaines_analyses', []))}")
            print(f"   â±ï¸ Temps: {processing_time}s")
            print(f"   ğŸ“Š Tokens: {total_tokens:,}")
            print(f"   ğŸ’° CoÃ»t: ${total_cost}")
            
            return matching_result
            
        except Exception as e:
            print(f"âŒ Erreur analyse matching: {e}", flush=True)
            import traceback
            print(traceback.format_exc(), flush=True)
            return {
                'score_matching': 0,
                'domaines_analyses': [],
                'synthese_matching': f'Erreur lors de l\'analyse: {str(e)}'
            }
    
    def enrich_cv_with_prompt(
        self, 
        parsed_cv: Dict[str, Any], 
        jd_text: str, 
        language: str = "French",
        matching_analysis: Dict[str, Any] = None  # âœ… FIX: Nouveau paramÃ¨tre pour rÃ©utiliser le matching
    ) -> Dict[str, Any]:
        """
        Enrichir le CV avec l'IA
        
        Args:
            parsed_cv: CV parsÃ©
            jd_text: Job Description
            language: Langue cible (French/English)
            matching_analysis: RÃ©sultat optionnel du matching prÃ©alable (Step 1)
                              Si fourni, rÃ©utilise le score au lieu de le recalculer
        
        Returns:
            CV enrichi avec tous les champs nÃ©cessaires
        """
        import time
        
        # âš ï¸ CRITICIAL: DÃ©terminer si on rÃ©utilise le scoring du Step 1
        reuse_scoring = matching_analysis is not None
        
        print(f"âœ¨ Enrichissement du CV avec l'IA...", flush=True)
        print(f"   Langue cible: {language}", flush=True)
        print(f"   Mode: {'RÃ©utilisation scoring Step 1' if reuse_scoring else 'Scoring complet'}", flush=True)
        
        # â±ï¸ DÃ©marrer le chronomÃ¨tre
        start_time = time.time()
        
        try:
            client = self._get_anthropic_client()
            
            # Reconstruire le CV en texte pour le prompt
            cv_text = f"""
PROFIL: {parsed_cv.get('profil_resume', '')}

TITRE: {parsed_cv.get('titre_professionnel', '')}

COMPÃ‰TENCES:
{chr(10).join(['- ' + comp for comp in parsed_cv.get('competences', [])])}

EXPÃ‰RIENCES:
"""
            for exp in parsed_cv.get('experiences', []):
                cv_text += f"\n{exp.get('periode', '')} | {exp.get('entreprise', '')} | {exp.get('poste', '')}\n"
                for resp in exp.get('responsabilites', []):
                    cv_text += f"  - {resp}\n"
            
            cv_text += "\nFORMATION:\n"
            for form in parsed_cv.get('formation', []):
                cv_text += f"- {form.get('diplome', '')} | {form.get('institution', '')} | {form.get('annee', '')}\n"
        
            # PROMPT ULTRA-RENFORCÃ‰ POUR COHÃ‰RENCE ABSOLUE
            language_instruction = f"""
ğŸš¨ RÃˆGLE ABSOLUE - LANGUE {language.upper()} ğŸš¨
â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”

âš ï¸ INSTRUCTION CRITIQUE - LANGUE OBLIGATOIRE: {language.upper()}

Tu DOIS gÃ©nÃ©rer 100% du contenu en {language} - AUCUNE EXCEPTION:
âœ“ Le TITRE PROFESSIONNEL en {language}
âœ“ Le PROFIL ENRICHI en {language}
âœ“ TOUTES les COMPÃ‰TENCES en {language}
âœ“ TOUTES les EXPÃ‰RIENCES en {language}
âœ“ TOUS les noms de catÃ©gories en {language}
âœ“ TOUTES les descriptions en {language}
âœ“ TOUS les mots-clÃ©s en {language}

ğŸ”´ SI {language} = "French":
- Utilise: "Analyste", "Gestion", "Configuration", "DÃ©veloppement", "Senior"
- PAS: "Analyst", "Management", "Development"
- Exemple titre: "Analyste QA Senior" âœ“ (PAS "Senior QA Analyst" âœ—)
- Exemple description: "Configuration de SharePoint incluant gestion..."
- Style: FranÃ§ais professionnel standard

ğŸ”´ SI {language} = "English":
- Utilise: "Analyst", "Management", "Configuration", "Development", "Senior"
- PAS: "Analyste", "Gestion", "DÃ©veloppement"
- Exemple titre: "Senior QA Analyst" âœ“ (PAS "Analyste QA Senior" âœ—)
- Exemple description: "SharePoint configuration including management..."
- Style: Professional English standard

IMPORTANT TITRE PROFESSIONNEL:
- Adapte le titre Ã  la Job Description
- Le titre doit Ãªtre COURT (3-5 mots maximum)
- Le titre doit Ãªtre en {language} - VÃ‰RIFIE 2 FOIS
- Si langue = French: ordre franÃ§ais (ex: "Analyste Configuration SharePoint")
- Si langue = English: ordre anglais (ex: "SharePoint Configuration Analyst")

VÃ‰RIFICATION FINALE OBLIGATOIRE:
Avant de rÃ©pondre, relis TOUT ton JSON et confirme que:
1. Le titre_professionnel_enrichi est en {language} âœ“
2. Le profil_enrichi est en {language} âœ“
3. Toutes les catÃ©gories de compÃ©tences sont en {language} âœ“
4. Toutes les descriptions sont en {language} âœ“
5. Les responsabilitÃ©s des expÃ©riences sont en {language} âœ“

Si UNE SEULE phrase n'est pas en {language} â†’ RECOMMENCE TOUT.

â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”â”

ğŸ¯ RÃ”LE CRITIQUE - TU ES UN RECRUTEUR SENIOR PROFESSIONNEL:
- Tu as 15+ ans d'expÃ©rience en recrutement technique
- Tu travailles pour le CLIENT (l'entreprise qui recrute)
- Ta mission: Ã©valuer si le CANDIDAT correspond EXACTEMENT aux besoins du CLIENT
- Tu dois Ãªtre OBJECTIF, RIGOUREUX et REPRODUCTIBLE dans ton Ã©valuation
- Ton scoring doit Ãªtre IDENTIQUE si tu analyses le mÃªme CV/JD plusieurs fois
- Tu notes comme un examinateur professionnel, pas comme un vendeur
"""
            
            # âœ… FIX: Choisir le prompt selon si on rÃ©utilise le matching ou non
            if reuse_scoring:
                # ============================================
                # VERSION SIMPLIFIÃ‰E - Matching dÃ©jÃ  fait au Step 1
                # ============================================
                prompt = f"""Voici la job description et le CV actuel ci-dessous.

ğŸ”¹ AmÃ©liore le CV pour qu'il soit parfaitement alignÃ© avec la job description tout en gardant le format d'origine (titres, mise en page, structure, ton professionnel).
{language_instruction}

âš ï¸ IMPORTANT: L'analyse de matching a DÃ‰JÃ€ Ã©tÃ© faite. Tu dois UNIQUEMENT faire l'ENRICHISSEMENT du contenu.

Fais :

1. Une version rÃ©Ã©crite et enrichie du CV

2a. TITRE: TITRE COURT adaptÃ© Ã  la JD en {language} (3-5 mots max)

2b. PROFIL exceptionnel : Ã©cris un paragraphe NARRATIF fluide (pas de liste), 5-6 lignes avec progression logique.

2c. GRAS ULTRA-SÃ‰LECTIF : identifie UNIQUEMENT 3-5 technologies CRITIQUES.

3. IntÃ¨gre naturellement les mots-clÃ©s techniques de la JD
4. Ajuste les intitulÃ©s pour que le profil paraisse livrable immÃ©diatement
5. N'invente rien â€” reformule uniquement les Ã©lÃ©ments prÃ©sents
6. EXPÃ‰RIENCES : bullets courts (1 ligne max), maximum 5-6 bullets par expÃ©rience

RÃ©ponds en JSON STRICT (sans markdown) avec cette structure:
{{
  "titre_professionnel_enrichi": "TITRE COURT en {language} (3-5 mots max)",
  
  "profil_enrichi": "Profil NARRATIF 5-6 lignes en {language} avec **3-5 technologies clÃ©s** en gras",
  
  "mots_cles_a_mettre_en_gras": ["Liste 15-20 TECHNOLOGIES de la JD - PAS de verbes gÃ©nÃ©riques"],
  
  "competences_enrichies": {{
    "Nom CatÃ©gorie 1 (3-6 mots max)": [
      "**Technologie principale** : description en 2-3 lignes (MAXIMUM 100-150 caractÃ¨res) incluant contexte, outils associÃ©s (**outil1**, **outil2**) et rÃ©sultats. Style concis et percutant.",
      "**Autre technologie** : description COURTE avec contexte + outils (**tech1**, **tech2**) + impact. 2-3 technologies en **gras** par compÃ©tence."
    ],
    "Nom CatÃ©gorie 2": [
      "CompÃ©tence concise..."
    ]
  }},
  
  RÃˆGLES ULTRA-CRITIQUES pour les compÃ©tences (NON-NÃ‰GOCIABLE):
  - Noms de catÃ©gories COURTS (3-6 mots max)
  - 5-6 catÃ©gories ADAPTÃ‰ES Ã  la JD
  - Chaque catÃ©gorie: 3-5 compÃ©tences MAXIMUM
  - CHAQUE compÃ©tence : 2-3 LIGNES MAXIMUM (100-150 caractÃ¨res) - NE PAS DÃ‰PASSER
  - Format: "**Technologie** : description concise avec outils (**outil1**, **outil2**) + rÃ©sultats"
  - 2-3 technologies en **gras** par compÃ©tence (PAS PLUS)
  - Descriptions CONCISES, CLAIRES et PROFESSIONNELLES
  - PrivilÃ©gier CLARTÃ‰ et CONCISION sur la longueur
  
  "experiences_enrichies": [
    {{
      "periode": "2020-2023",
      "entreprise": "Nom entreprise",
      "poste": "Titre reformulÃ© selon JD",
      "responsabilites": [
        "Configuration **Open edX** incluant structuration et intÃ©gration avec **SharePoint** pour gestion contenus",
        "Automatisation processus documentaires via **Power Automate** et **Teams** pour amÃ©liorer efficacitÃ©"
      ],
      "environment": "**Open edX**, **SharePoint**, **Microsoft 365**, Teams, Power Automate, OneDrive, SQL"
    }}
  ]
}}

FORMAT OBLIGATOIRE (COPIER format compÃ©tences):
- ResponsabilitÃ©s: Technologies **isolÃ©es** dans texte normal (ex: "Configuration **Tech1** incluant **Tech2** pour rÃ©sultats")
- Environnement: Liste virgules avec 3-5 technologies **critiques** en gras, autres sans
- JAMAIS phrases entiÃ¨res en gras
- Maximum 2-3 mots entre **astÃ©risques**

---

JOB DESCRIPTION:
{jd_text}

---

CV ACTUEL:
{cv_text}

---

IMPORTANT FINAL - RÃˆGLES JSON STRICTES:
- GÃ©nÃ¨re UNIQUEMENT du JSON valide
- PAS de commentaires (// ou /* */)
- PAS de virgules finales (trailing commas)
- PAS de markdown (```json ou ```)
- TOUS les strings doivent utiliser des guillemets doubles ""
- VÃ©rifie que TOUTES les accolades et crochets sont fermÃ©s
- Si tu hÃ©sites sur un champ, mets une valeur par dÃ©faut plutÃ´t qu'une erreur

RÃ©ponds UNIQUEMENT avec du JSON pur, sans rien d'autre avant ou aprÃ¨s."""

            else:
                # ============================================
                # VERSION COMPLÃˆTE - Mode legacy/fallback avec matching inclus
                # ============================================
                prompt = f"""Voici la job description et le CV actuel ci-dessous.

ğŸ”¹ AmÃ©liore le CV pour qu'il soit parfaitement alignÃ© avec la job description tout en gardant le format d'origine (titres, mise en page, structure, ton professionnel).
{language_instruction}

ğŸ¯ ANALYSE DE MATCHING PONDÃ‰RÃ‰E (ULTRA-CRITIQUE - COHÃ‰RENCE ABSOLUE REQUISE):

âš ï¸ PRINCIPE FONDAMENTAL DE COHÃ‰RENCE - MÃ‰THODOLOGIE STRICTE:
- Tu es un SYSTÃˆME D'Ã‰VALUATION AUTOMATISÃ‰, pas un humain
- Pour le MÃŠME CV et la MÃŠME JD â†’ EXACTEMENT le mÃªme score Ã  chaque fois
- Utilise une grille d'Ã©valuation MATHÃ‰MATIQUE et REPRODUCTIBLE
- Agis comme un ALGORITHME, pas comme un recruteur subjectif
- Chaque critÃ¨re suit des rÃ¨gles BINAIRES strictes (oui/non, prÃ©sent/absent)
- Tu DOIS pouvoir justifier CHAQUE point attribuÃ© avec des FAITS du CV
- Si tu hÃ©sites entre 2 scores â†’ prends le PLUS BAS (principe de strictness)

ğŸ”´ RÃˆGLE D'OR - SCORE GLOBAL = SOMME DOMAINES:
- Le score_matching FINAL = somme EXACTE de tous les scores de domaines
- VÃ‰RIFIE 3 FOIS avant de rÃ©pondre: somme des scores = score_matching
- Si tu calcules 37/100 en sommant les domaines â†’ score_matching DOIT Ãªtre 37
- NE JAMAIS inventer un score global diffÃ©rent de la somme calculÃ©e

Ã‰TAPE 1 - IDENTIFIER 5-8 DOMAINES CRITIQUES (MÃ‰THODE ALGORITHIMQUE):

ğŸ“‹ PROCESSUS AUTOMATIQUE D'IDENTIFICATION:
1. Scan complet de la JD - repÃ©rer TOUS les mots techniques
2. Compter la frÃ©quence EXACTE de chaque technologie/compÃ©tence
3. CrÃ©er une liste de domaines par ordre d'importance
4. Appliquer la formule de pondÃ©ration ci-dessous

ğŸ“Š FORMULE DE PONDÃ‰RATION MATHÃ‰MATIQUE:
Pour chaque domaine, calcule son poids avec:

Poids_Base = (Nombre_mentions / Total_mentions_techniques) Ã— 100

Bonus:
- +20% si c'est le TITRE du poste (ex: ".NET Developer" â†’ Stack .NET = +20%)
- +15% si mots "Required", "Must have", "Essential", "Critical"
- +10% si mentionnÃ© dans les 3 premiÃ¨res lignes de la JD
- +5% par occurrence au-delÃ  de 3 mentions

Poids_Final = min(Poids_Base + Bonus, 50%)  â† Aucun domaine ne peut dÃ©passer 50%

RÃˆGLES STRICTES DE PONDÃ‰RATION:
- Stack technique principal (dans titre ou 5+ mentions): 30-50%
- CompÃ©tences techniques secondaires (3-4 mentions): 15-25%
- CompÃ©tences techniques tertiaires (1-2 mentions): 5-15%
- Soft skills/Leadership: MAXIMUM 10% (sauf si poste management)
- TOTAL des poids = EXACTEMENT 100% (vÃ©rifie avec calculatrice)
- Si total â‰  100%, ajuste proportionnellement tous les poids

Ã‰TAPE 2 - SCORER CHAQUE DOMAINE (ALGORITHME DE NOTATION STRICT):

ğŸ¤– SYSTÃˆME DE NOTATION AUTOMATISÃ‰ - APPLIQUE CES RÃˆGLES EXACTEMENT:

POUR CHAQUE DOMAINE, COMPTE:
a) Nombre de mentions EXACTES de la technologie dans le CV
b) Nombre de projets/expÃ©riences utilisant cette technologie  
c) DurÃ©e totale d'utilisation (annÃ©es)
d) Niveau dÃ©montrÃ© (junior/intermÃ©diaire/senior)

ğŸ“ FORMULE MATHÃ‰MATIQUE DE SCORING:

Ã‰tape 2A - Score Brut (0-100%):
â€¢ 0% : ZÃ‰RO mention de la techno dans le CV, stack incompatible
â€¢ 10% : Technologie proche mentionnÃ©e (PostgreSQL pour SQL Server)
â€¢ 25% : 1 mention + aucune expÃ©rience pratique (formation seulement)
â€¢ 40% : 1-2 mentions + 1 projet + <1 an d'expÃ©rience
â€¢ 60% : 3-4 mentions + 2 projets + 1-2 ans d'expÃ©rience
â€¢ 80% : 5+ mentions + 3+ projets + 3+ ans d'expÃ©rience
â€¢ 100% : 7+ mentions + expertise dÃ©montrÃ©e + senior confirmÃ©

Ã‰tape 2B - Ajustements OBLIGATOIRES:
â€¢ Si stack incompatible (Java vs .NET) â†’ Score = 0% (NON-NÃ‰GOCIABLE)
â€¢ Si technologie absente du CV â†’ Score = 0% (NON-NÃ‰GOCIABLE)
â€¢ Si aucune expÃ©rience pratique prouvÃ©e â†’ Score MAX = 30%
â€¢ Si expÃ©rience <1 an â†’ Score MAX = 50%
â€¢ Si niveau junior Ã©vident â†’ Score MAX = 60%

Ã‰tape 2C - Calcul Final:
Score_Domaine = (Score_Brut Ã— Poids_Domaine) / 100

EXEMPLE DÃ‰TAILLÃ‰:
Domaine: ".NET Development" - Poids: 40%
CV candidat: AUCUNE mention .NET, seulement Java
â†’ Score_Brut = 0%
â†’ Score_Domaine = (0 Ã— 40) / 100 = 0 points
â†’ Commentaire: "âŒ Stack incompatible - profil Java exclusif"

ğŸ”´ VÃ‰RIFICATION FINALE OBLIGATOIRE:
Somme_Scores = Î£(tous les Score_Domaine)
Si Somme_Scores â‰  score_matching â†’ ERREUR CRITIQUE â†’ RECALCULE

Ã‰TAPE 3 - COMMENTAIRE PAR DOMAINE (30-50 mots):
- Utilise âŒ (0-30%), âš ï¸ (30-70%), âœ… (70-100%)
- Sois FACTUEL et OBJECTIF dans tes commentaires
- Base-toi UNIQUEMENT sur les FAITS prÃ©sents dans le CV
- Ne fais PAS d'hypothÃ¨ses optimistes

EXEMPLE:
JD demande: ".NET, C#, Azure, SQL Server"
Candidat a: "Java, AWS, PostgreSQL"

RÃ‰SULTAT:
{{
  "domaines_analyses": [
    {{
      "domaine": "Stack .NET (C#, ASP.NET Core, Entity Framework)",
      "poids": 40,
      "score": 0,
      "score_max": 40,
      "commentaire": "âŒ Aucune expÃ©rience .NET/C#. Profil Java exclusivement - incompatibilitÃ© majeure sur stack principale.",
      "match": "incompatible"
    }},
    {{
      "domaine": "Cloud Microsoft Azure",
      "poids": 20,
      "score": 8,
      "score_max": 20,
      "commentaire": "âš ï¸ ExpÃ©rience AWS uniquement. CompÃ©tences cloud transfÃ©rables mais nÃ©cessite formation Azure.",
      "match": "partiel"
    }},
    {{
      "domaine": "SQL Server & T-SQL",
      "poids": 15,
      "score": 10,
      "score_max": 15,
      "commentaire": "âœ… MaÃ®trise PostgreSQL et MySQL - compÃ©tences SQL transfÃ©rables Ã  SQL Server.",
      "match": "bon"
    }}
  ],
  "score_matching": 45,
  "synthese_matching": "PARTIAL MATCH (45/100) - This Java senior profile presents significant challenges for a .NET-focused role, though some transferable competencies exist.

KEY STRENGTHS: The candidate brings 8+ years of enterprise software development experience with proven expertise in cloud platforms (AWS/Azure) and database technologies (PostgreSQL, MySQL). Their experience leading technical teams and architecting scalable solutions demonstrates strong senior-level capabilities. The containerization skills (Docker, Kubernetes) mentioned in their current role are highly relevant.

PARTIAL MATCHES: While the candidate's SQL database experience is strong and transferable to SQL Server, their cloud platform knowledge (AWS/Azure fundamentals) provides a foundation that could accelerate learning of Azure-specific services required for this role. Their experience with agile methodologies and team leadership aligns well with the position's requirements.

CRITICAL GAPS: The most significant concern is the complete absence of .NET stack experience (C#, ASP.NET, Entity Framework), which represents 40% of the role's core requirements (0/40 points). The candidate would require substantial retraining on the entire Microsoft technology stack. Additionally, there's no evidence of Azure-specific service experience (Azure Functions, Service Bus, etc.) beyond basic cloud concepts.

RECOMMENDATION: This profile requires major reconversion and is NOT recommended for immediate placement. Consider only if: (1) the client accepts a 3-6 month ramp-up period, (2) candidate demonstrates strong motivation to transition to .NET, and (3) budget allows for extensive training investment. For urgent needs, seek candidates with existing .NET experience."
}}

Fais :

1. ANALYSE PONDÃ‰RÃ‰E OBLIGATOIRE (voir ci-dessus)
2. Une version rÃ©Ã©crite et enrichie du CV

2b. PROFIL exceptionnel : Ã©cris un paragraphe NARRATIF fluide (pas de liste), 5-6 lignes avec progression logique.

2c. GRAS ULTRA-SÃ‰LECTIF : identifie UNIQUEMENT 3-5 technologies CRITIQUES.

3. IntÃ¨gre naturellement les mots-clÃ©s techniques de la JD
4. Ajuste les intitulÃ©s pour que le profil paraisse livrable immÃ©diatement
5. N'invente rien â€” reformule uniquement les Ã©lÃ©ments prÃ©sents
6. EXPÃ‰RIENCES : bullets courts (1 ligne max), maximum 5-6 bullets par expÃ©rience

RÃ©ponds en JSON STRICT (sans markdown) avec cette structure:
{{
  "domaines_analyses": [
    {{
      "domaine": "Nom domaine technique/fonctionnel (ex: Stack .NET, Cloud Azure)",
      "poids": 40,
      "score": 15,
      "score_max": 40,
      "commentaire": "Explication 30-50 mots avec âŒ/âš ï¸/âœ…",
      "match": "incompatible|partiel|bon|excellent"
    }}
  ],
  "score_matching": 45,
  "synthese_matching": "CONCISE PROFESSIONAL SUMMARY (1 paragraph, 80-120 words):
  
  Write a single comprehensive paragraph that includes:
  - Match level (Excellent/Strong/Good/Partial/Weak) with the score (X/100)
  - Candidate's years of experience and seniority level
  - Top 2-3 strongest domains that align perfectly with requirements
  - 1-2 areas that are partial matches or transferable skills
  - 1-2 critical gaps if any
  - Brief recommendation (Recommend/Conditional/Not Recommend)
  
  Keep it analytical and professional. Use concrete examples from the CV. Be honest about both strengths and weaknesses. Make it scannable for busy recruiters. ALWAYS WRITE IN ENGLISH.",
  
  "titre_professionnel_enrichi": "TITRE COURT en {language} (3-5 mots max)",
  
  "profil_enrichi": "Profil NARRATIF 5-6 lignes en {language} avec **3-5 technologies clÃ©s** en gras",
  
  "mots_cles_a_mettre_en_gras": ["Liste 15-20 TECHNOLOGIES de la JD - PAS de verbes gÃ©nÃ©riques"],
  
  "competences_enrichies": {{
    "Nom CatÃ©gorie 1 (3-6 mots max)": [
      "**Technologie principale** : description en 2-3 lignes (MAXIMUM 100-150 caractÃ¨res) incluant contexte, outils associÃ©s (**outil1**, **outil2**) et rÃ©sultats. Style concis et percutant.",
      "**Autre technologie** : description COURTE avec contexte + outils (**tech1**, **tech2**) + impact. 2-3 technologies en **gras** par compÃ©tence."
    ],
    "Nom CatÃ©gorie 2": [
      "CompÃ©tence concise..."
    ]
  }},
  
  RÃˆGLES ULTRA-CRITIQUES pour les compÃ©tences (NON-NÃ‰GOCIABLE):
  - Noms de catÃ©gories COURTS (3-6 mots max)
  - 5-6 catÃ©gories ADAPTÃ‰ES Ã  la JD
  - Chaque catÃ©gorie: 3-5 compÃ©tences MAXIMUM
  - CHAQUE compÃ©tence : 2-3 LIGNES MAXIMUM (100-150 caractÃ¨res) - NE PAS DÃ‰PASSER
  - Format: "**Technologie** : description concise avec outils (**outil1**, **outil2**) + rÃ©sultats"
  - 2-3 technologies en **gras** par compÃ©tence (PAS PLUS)
  - Descriptions CONCISES, CLAIRES et PROFESSIONNELLES
  - PrivilÃ©gier CLARTÃ‰ et CONCISION sur la longueur
  
  "experiences_enrichies": [
    {{
      "periode": "2020-2023",
      "entreprise": "Nom entreprise",
      "poste": "Titre reformulÃ© selon JD",
      "responsabilites": [
        "Configuration **Open edX** incluant structuration et intÃ©gration avec **SharePoint** pour gestion contenus",
        "Automatisation processus documentaires via **Power Automate** et **Teams** pour amÃ©liorer efficacitÃ©"
      ],
      "environment": "**Open edX**, **SharePoint**, **Microsoft 365**, Teams, Power Automate, OneDrive, SQL"
    }}
  ],
  
  FORMAT OBLIGATOIRE (COPIER format compÃ©tences):
  - ResponsabilitÃ©s: Technologies **isolÃ©es** dans texte normal (ex: "Configuration **Tech1** incluant **Tech2** pour rÃ©sultats")
  - Environnement: Liste virgules avec 3-5 technologies **critiques** en gras, autres sans
  - JAMAIS phrases entiÃ¨res en gras
  - Maximum 2-3 mots entre **astÃ©risques**
  
  "score_matching": 45,
  "points_forts": ["ALWAYS in English: key strength 1", "ALWAYS in English: key strength 2"]
}}

ğŸŒ CRITICAL LANGUAGE REQUIREMENT:
- 'domaines_analyses' (domain names AND comments) MUST ALWAYS be in ENGLISH
- 'synthese_matching' MUST ALWAYS be in ENGLISH  
- 'points_forts' MUST ALWAYS be in ENGLISH
- Example domain: "SQL Data Extraction and Manipulation" NOT "Extraction de donnÃ©es SQL"
- Example comment: "âŒ No demonstrated experience in SQL data extraction..." NOT "âŒ Aucune expÃ©rience..."
- Example synthesis: "Java senior profile unsuitable for .NET position..." NOT "Profil Java senior inadaptÃ©..."

CRITICAL SCORING RULES:
- 'domaines_analyses' MUST be completed with 5-8 domains totaling EXACTLY 100%
- BE STRICT on scoring - don't give points if candidate lacks the skill
- If stack mismatch (Java vs .NET), give 0 points, not 40-50

ğŸ”´ğŸ”´ğŸ”´ VÃ‰RIFICATION FINALE AVANT RÃ‰PONSE (NON-NÃ‰GOCIABLE) ğŸ”´ğŸ”´ğŸ”´

AVANT de gÃ©nÃ©rer ta rÃ©ponse JSON, tu DOIS:

1ï¸âƒ£ CALCULER LA SOMME:
   Somme = domaine1.score + domaine2.score + domaine3.score + ... + domaineN.score
   
2ï¸âƒ£ VÃ‰RIFIER:
   Si Somme â‰  score_matching â†’ ERREUR â†’ RECALCULE TOUT
   
3ï¸âƒ£ VÃ‰RIFIER LES POIDS:
   Somme_Poids = domaine1.poids + domaine2.poids + ... + domaineN.poids
   Si Somme_Poids â‰  100 â†’ ERREUR â†’ RECALCULE TOUT
   
4ï¸âƒ£ DOUBLE-CHECK:
   Pour chaque domaine: vÃ©rifie que score â‰¤ score_max
   Pour chaque domaine: vÃ©rifie que score_max = poids

EXEMPLE DE VÃ‰RIFICATION:
Domaine 1: Stack .NET (40%) â†’ 0/40 points
Domaine 2: Cloud Azure (20%) â†’ 8/20 points  
Domaine 3: SQL Server (15%) â†’ 10/15 points
Domaine 4: DevOps (15%) â†’ 5/15 points
Domaine 5: Agile (10%) â†’ 7/10 points

VÃ©rification poids: 40+20+15+15+10 = 100 âœ…
VÃ©rification score: 0+8+10+5+7 = 30 âœ…
Donc: score_matching = 30 âœ…

Si tu trouves une incohÃ©rence â†’ RECALCULE TOUT depuis le dÃ©but

---

JOB DESCRIPTION:
{jd_text}

---

CV ACTUEL:
{cv_text}

---

IMPORTANT FINAL - RÃˆGLES JSON STRICTES:
- GÃ©nÃ¨re UNIQUEMENT du JSON valide
- PAS de commentaires (// ou /* */)
- PAS de virgules finales (trailing commas)
- PAS de markdown (```json ou ```)
- TOUS les strings doivent utiliser des guillemets doubles ""
- VÃ©rifie que TOUTES les accolades et crochets sont fermÃ©s
- Si tu hÃ©sites sur un champ, mets une valeur par dÃ©faut plutÃ´t qu'une erreur

RÃ©ponds UNIQUEMENT avec du JSON pur, sans rien d'autre avant ou aprÃ¨s."""

            print(f">>> Calling Claude API for enrichment with timeout=300s...", flush=True)
            response = client.messages.create(
                model="claude-sonnet-4-5-20250929",
                max_tokens=8000,
                timeout=300.0,  # 5 minutes max
                messages=[{"role": "user", "content": prompt}]
            )
            print(f">>> Enrichment API call completed successfully", flush=True)
            
            # ğŸ“Š Capturer les mÃ©tadonnÃ©es API
            input_tokens = response.usage.input_tokens if hasattr(response, 'usage') else 0
            output_tokens = response.usage.output_tokens if hasattr(response, 'usage') else 0
            total_tokens = input_tokens + output_tokens
            
        except Exception as e:
            print(f">>> ERROR calling anthropic for enrichment: {repr(e)}", flush=True)
            import traceback
            print(f">>> FULL TRACEBACK:\n{traceback.format_exc()}", flush=True)
            return {}
        
        print(f">>> API Response received, extracting text...", flush=True)
        response_text = response.content[0].text.strip()
        print(f">>> Response length: {len(response_text)} characters", flush=True)
        print(f">>> Response preview (first 500 chars):\n{response_text[:500]}", flush=True)
        
        # Nettoyer JSON
        if response_text.startswith('```json'):
            response_text = response_text[7:]
        if response_text.startswith('```'):
            response_text = response_text[3:]
        if response_text.endswith('```'):
            response_text = response_text[:-3]
        response_text = response_text.strip()
        
        print(f">>> Attempting to parse JSON...", flush=True)
        
        # ğŸ”§ NOUVEAU: Tentative de parsing avec retry et correction
        enriched = None
        max_retries = 3
        
        for attempt in range(max_retries):
            try:
                if attempt == 0:
                    # PremiÃ¨re tentative: parsing direct
                    enriched = json.loads(response_text)
                    print(f">>> JSON parsed successfully on first attempt!", flush=True)
                    break
                else:
                    # Tentatives suivantes: demander Ã  Claude de corriger le JSON
                    print(f">>> Retry {attempt}/{max_retries-1}: Asking Claude to fix JSON...", flush=True)
                    
                    fix_prompt = f"""The following JSON is malformed. Please fix it and return ONLY the corrected JSON without any explanation or markdown:

{response_text}

Return the corrected JSON directly:"""
                    
                    fix_response = client.messages.create(
                        model="claude-sonnet-4-5-20250929",
                        max_tokens=8000,
                        timeout=60.0,
                        messages=[{"role": "user", "content": fix_prompt}]
                    )
                    
                    fixed_text = fix_response.content[0].text.strip()
                    # Nettoyer le JSON corrigÃ©
                    if fixed_text.startswith('```json'):
                        fixed_text = fixed_text[7:]
                    if fixed_text.startswith('```'):
                        fixed_text = fixed_text[3:]
                    if fixed_text.endswith('```'):
                        fixed_text = fixed_text[:-3]
                    fixed_text = fixed_text.strip()
                    
                    enriched = json.loads(fixed_text)
                    print(f">>> JSON successfully fixed and parsed on attempt {attempt}!", flush=True)
                    break
                    
            except json.JSONDecodeError as e:
                print(f"âš ï¸ Erreur JSON (attempt {attempt + 1}/{max_retries}): {e}", flush=True)
                if attempt == 0:
                    print(f">>> JSON Error position: {e.pos}", flush=True)
                    print(f">>> Problematic section: {response_text[max(0, e.pos-100):e.pos+100]}", flush=True)
                
                if attempt == max_retries - 1:
                    # Dernier essai Ã©chouÃ©: retourner dict vide
                    print(f">>> All parsing attempts failed. Returning empty dict.", flush=True)
                    print(f">>> Full response text:\n{response_text}", flush=True)
                    return {}
                else:
                    # Continuer au prochain retry
                    continue
        
        if enriched is None:
            print(f">>> ERROR: enriched is None after all retries", flush=True)
            return {}
        
        print(f">>> Keys in enriched: {list(enriched.keys())}", flush=True)
        
        # â±ï¸ Calculer le temps de traitement
        processing_time = round(time.time() - start_time, 2)
        
        # ğŸ’° Calculer le coÃ»t (prix Claude Sonnet 4.5: $3/MTok input, $15/MTok output)
        cost_input = (input_tokens / 1_000_000) * 3.0
        cost_output = (output_tokens / 1_000_000) * 15.0
        total_cost = round(cost_input + cost_output, 4)
        
        # ğŸ“ˆ Ajouter les mÃ©tadonnÃ©es dans le rÃ©sultat
        enriched['_metadata'] = {
            'processing_time_seconds': processing_time,
            'input_tokens': input_tokens,
            'output_tokens': output_tokens,
            'total_tokens': total_tokens,
            'estimated_cost_usd': total_cost
        }
        
        print(f"âœ… Enrichissement rÃ©ussi!")
        
        # âœ… FIX: Si on rÃ©utilise le matching, merger les rÃ©sultats
        if reuse_scoring and matching_analysis:
            print(f"   Mode: RÃ©utilisation du matching du Step 1", flush=True)
            # RÃ©cupÃ©rer les rÃ©sultats du Step 1
            enriched['score_matching'] = matching_analysis.get('score_matching', 0)
            enriched['domaines_analyses'] = matching_analysis.get('domaines_analyses', [])
            enriched['synthese_matching'] = matching_analysis.get('synthese_matching', '')
            enriched['points_forts'] = matching_analysis.get('points_forts', [])
            print(f"   Score rÃ©utilisÃ©: {enriched['score_matching']}/100")
            print(f"   Domaines rÃ©utilisÃ©s: {len(enriched['domaines_analyses'])}")
        else:
            print(f"   Mode: Calcul complet du matching", flush=True)
            print(f"   Score matching: {enriched.get('score_matching', 0)}/100")
            print(f"   Domaines analysÃ©s: {len(enriched.get('domaines_analyses', []))}")
        print(f"   Mots-clÃ©s en gras: {len(enriched.get('mots_cles_a_mettre_en_gras', []))}")
        print(f"   â±ï¸ Temps de traitement: {processing_time}s")
        print(f"   ğŸ“Š Tokens: {total_tokens:,} ({input_tokens:,} in + {output_tokens:,} out)")
        print(f"   ğŸ’° CoÃ»t estimÃ©: ${total_cost}")
        
        if enriched.get('domaines_analyses'):
            print(f"\n   ğŸ“Š DÃ©tail scoring:")
            for domaine in enriched['domaines_analyses']:
                emoji = domaine.get('match', '')
                if emoji == 'incompatible':
                    emoji = 'âŒ'
                elif emoji == 'partiel':
                    emoji = 'âš ï¸'
                elif emoji in ['bon', 'excellent']:
                    emoji = 'âœ…'
                print(f"      {emoji} {domaine.get('domaine', 'N/A')}: {domaine.get('score', 0)}/{domaine.get('score_max', 0)} ({domaine.get('poids', 0)}%)")
        
        # DEBUG: Afficher une responsabilitÃ© pour voir le format
        if enriched.get('experiences_enrichies'):
            first_exp = enriched['experiences_enrichies'][0]
            if first_exp.get('responsabilites'):
                print(f"\nğŸ” DEBUG - PremiÃ¨re responsabilitÃ© :")
                print(f"   {first_exp['responsabilites'][0]}")
            if first_exp.get('environment'):
                print(f"\nğŸ” DEBUG - Environnement :")
                print(f"   {first_exp['environment']}")
        
        # ğŸš¨ VÃ©rification critique: le dict ne doit pas Ãªtre vide
        if not enriched:
            print(f">>> WARNING: enriched dict is EMPTY!", flush=True)
            return {}
        
        # VÃ©rifier les clÃ©s essentielles
        required_keys = ['score_matching', 'domaines_analyses', 'profil_enrichi']
        missing_keys = [k for k in required_keys if k not in enriched]
        if missing_keys:
            print(f">>> WARNING: Missing critical keys: {missing_keys}", flush=True)
            print(f">>> Available keys: {list(enriched.keys())}", flush=True)
        
        return enriched

    # ========================================
    # MODULE 4 : MAPPING TMC + RICHTEXT
    # ========================================
    
    def mdbold_to_richtext(self, s: str) -> RichText:
        """Convertit les **bold** markdown en RichText propre sans cascade de gras."""
        import re
        rt = RichText()
        pattern = re.compile(r'\*\*(.*?)\*\*')
        last_end = 0

        # Ajouter le texte avant chaque bloc en gras
        for match in pattern.finditer(s):
            if match.start() > last_end:
                rt.add(s[last_end:match.start()], bold=False, font='Arial')
            # Le texte entre **...** est en gras
            rt.add(match.group(1), bold=True, font='Arial')
            last_end = match.end()

        # Ajouter le texte aprÃ¨s le dernier bloc
        if last_end < len(s):
            rt.add(s[last_end:], bold=False, font='Arial')

        return rt

    def map_to_tmc_structure(self, parsed_cv: Dict[str, Any], enriched_cv: Dict[str, Any], template_lang: str = 'FR') -> Dict[str, Any]:
        """Mapper les donnÃ©es enrichies vers la structure TMC"""
        print("ğŸ—ºï¸  Mapping vers structure TMC...")
        
        # 1. PROFIL - Convertir en RichText pour supporter le gras (pas d'Ã©chappement)
        profil_brut = enriched_cv.get('profil_enrichi', parsed_cv.get('profil_resume', ''))
        profil = self.mdbold_to_richtext(profil_brut) if profil_brut else ''
        
        # 2. COMPÃ‰TENCES - FORMAT CATÃ‰GORISÃ‰ DÃ‰TAILLÃ‰
        competences_enrichies = enriched_cv.get('competences_enrichies', {})
        
        # Si competences_enrichies est un dict (nouveau format), l'utiliser directement
        if isinstance(competences_enrichies, dict):
            # Supprimer la clÃ© "NOTE" si prÃ©sente
            skills_categorized = {k: v for k, v in competences_enrichies.items() if k != 'NOTE' and isinstance(v, list)}
        else:
            # Fallback ancien format (liste simple)
            competences = competences_enrichies if isinstance(competences_enrichies, list) else parsed_cv.get('competences', [])
            skills_categorized = {
                'CompÃ©tences techniques': competences[:8] if len(competences) >= 8 else competences,
                'CompÃ©tences transversales': competences[8:12] if len(competences) > 8 else []
            }
            # Supprimer les catÃ©gories vides
            skills_categorized = {k: v for k, v in skills_categorized.items() if v}
        
        # ğŸ”¥ Transformation en RichText pour le formatage (pas d'Ã©chappement)
        skills_categorized_doc = []
        for cat, skills in skills_categorized.items():
            rt_cat = RichText()
            rt_cat.add(cat, bold=True)
            rt_skills = [self.mdbold_to_richtext(s) for s in skills]
            skills_categorized_doc.append((rt_cat, rt_skills))
        
        # 3. EXPÃ‰RIENCES - Texte simple pour les responsabilitÃ©s, RichText pour environnement
        experiences_enrichies = enriched_cv.get('experiences_enrichies', parsed_cv.get('experiences', []))
        work_experience = []
        
        for exp in experiences_enrichies:
            # GARDER les responsabilitÃ©s en TEXTE SIMPLE (pas RichText) - pas d'Ã©chappement
            responsabilites_text = [r for r in exp.get('responsabilites', [])]
            
            # Convertir l'environnement en RichText pour le gras - pas d'Ã©chappement
            environment_brut = exp.get('environment', '')
            environment_rt = self.mdbold_to_richtext(environment_brut) if environment_brut else ''
            
            work_exp = {
                'period': exp.get('periode', ''),
                'company': exp.get('entreprise', ''),
                'position': exp.get('poste', ''),
                'general_responsibilities': responsabilites_text,  # Texte simple
                'environment': environment_rt
            }
            work_experience.append(work_exp)
        
        # 4. FORMATION (avec dÃ©tails complets)
        formation = parsed_cv.get('formation', [])
        education = []
        for form in formation:
            education.append({
                'institution': form.get('institution', ''),
                'degree': form.get('diplome', ''),
                'graduation_year': form.get('annee', 'Date inconnue'),
                'country': form.get('pays', 'Canada'),
                'level': '',
                'title': form.get('diplome', '')
            })
        
        # 5. CERTIFICATIONS (avec mapping vers format template)
        certifications_raw = parsed_cv.get('certifications', [])
        certifications = []
        for cert in certifications_raw:
            certifications.append({
                'name': cert.get('nom', cert.get('name', '')),
                'institution': cert.get('organisme', cert.get('institution', '')),
                'year': str(cert.get('annee', cert.get('year', ''))),
                'country': cert.get('pays', cert.get('country', ''))
            })
        
        # 6. PROJETS
        projects = parsed_cv.get('projets', [])
        
        # 7. INFORMATIONS PERSONNELLES
        nom_complet = parsed_cv.get('nom_complet', '')
        
        # SÃ©parer prÃ©nom et nom
        parts = nom_complet.split() if nom_complet else []
        if len(parts) >= 2:
            first_name = parts[0]
            last_name = ' '.join(parts[1:])
        elif len(parts) == 1:
            first_name = parts[0]
            last_name = ''
        else:
            first_name = 'PrÃ©nom'
            last_name = 'Nom'
        
        titre_professionnel = enriched_cv.get('titre_professionnel_enrichi', parsed_cv.get('titre_professionnel', ''))
        lieu_residence = parsed_cv.get('lieu_residence', 'MontrÃ©al, Canada')
        langues_list = parsed_cv.get('langues', ['FranÃ§ais', 'Anglais'])
        
        # Traduire les langues selon le template
        if template_lang == 'FR':
            # Si template FR, traduire de l'anglais vers le franÃ§ais
            langue_map = {
                'English': 'Anglais',
                'French': 'FranÃ§ais',
                'Hebrew': 'HÃ©breu',
                'Russian': 'Russe',
                'Spanish': 'Espagnol',
                'German': 'Allemand',
                'Italian': 'Italien',
                'Portuguese': 'Portugais',
                'Chinese': 'Chinois',
                'Japanese': 'Japonais',
                'Arabic': 'Arabe'
            }
            langues_list = [langue_map.get(lang, lang) for lang in langues_list]
        
        langues = ', '.join(langues_list)
        
        context = {
            # Pour le header (minuscules) - PAS d'Ã©chappement
            'first_name': first_name,
            'last_name': last_name,
            'title': titre_professionnel,
            
            # Pour la page 1 (MAJUSCULES) - PAS d'Ã©chappement
            'FIRST_NAME': first_name.upper(),
            'LAST_NAME': last_name.upper(),
            'TITLE': titre_professionnel,
            'RESIDENCY': lieu_residence,
            'LANGUAGES': langues,
            
            # AUSSI en minuscules pour compatibilitÃ© template
            'residency': lieu_residence,
            'languages': langues,
            
            # Reste du CV
            'summary': profil,
            'skills_categorized': skills_categorized,
            'skills_categorized_doc': skills_categorized_doc,  # ğŸ”¥ Version RichText pour le template
            'work_experience': work_experience,
            'education': education,
            'projects': projects,
            'certifications': certifications
        }
        
        print(f"âœ… Mapping terminÃ©!")
        print(f"   Nom: [ANONYMIZED]")
        print(f"   Titre: {titre_professionnel}")
        print(f"   Langues: {langues}")
        print(f"   Profil: RichText gÃ©nÃ©rÃ©")
        total_competences = sum(len(v) for v in skills_categorized.values() if isinstance(v, list))
        print(f"   CatÃ©gories: {len(skills_categorized)}")
        print(f"   CompÃ©tences: {total_competences}")
        print(f"   ExpÃ©riences: {len(work_experience)}")
        
        return context

    # ========================================
    # MODULE 5 : GÃ‰NÃ‰RATION DOCX TMC
    # ========================================
    
    def find_template_file(self, template_name: str = "TMC_NA_template_FR.docx") -> str:
        """Recherche intelligente du template dans plusieurs emplacements possibles"""
        from pathlib import Path
        
        # Liste exhaustive des endroits possibles
        script_dir = Path(__file__).parent
        possible_paths = [
            Path(template_name),  # Current directory
            script_dir / template_name,  # Script directory
            script_dir.parent / "branding" / "templates" / template_name,  # ../../branding/templates/
            script_dir.parent.parent / "branding" / "templates" / template_name,  # ../../../branding/templates/
            Path.home() / template_name,  # Home directory
            Path.home() / "tmc-cv-optimizer" / "branding" / "templates" / template_name,  # Project in home
            Path("/app/branding/templates") / template_name,  # Render deployment path
            Path("/home/ubuntu/tmc-cv-optimizer/branding/templates") / template_name,  # Ubuntu deployment
        ]
        
        # Chercher dans les variables d'environnement aussi
        env_template_path = os.getenv("TMC_TEMPLATE_PATH")
        if env_template_path:
            possible_paths.insert(0, Path(env_template_path))
        
        print(f"   ğŸ” Recherche du template: {template_name}")
        
        for path in possible_paths:
            try:
                if path.exists() and path.is_file():
                    print(f"   âœ… Template trouvÃ©: {path.resolve()}")
                    return str(path.resolve())
            except (OSError, PermissionError) as e:
                # Ignorer silencieusement les erreurs de permissions
                continue
        
        # Si pas trouvÃ©, afficher tous les chemins essayÃ©s
        print(f"   âŒ Template introuvable: {template_name}")
        print(f"   Chemins testÃ©s:")
        for path in possible_paths:
            print(f"      - {path}")
        print(f"\n   ğŸ’¡ Astuce: DÃ©finir TMC_TEMPLATE_PATH pour spÃ©cifier un emplacement personnalisÃ©")
        raise FileNotFoundError(f"Template TMC introuvable: {template_name}")
    
    def generate_tmc_docx(self, context: Dict[str, Any], output_path: str, template_path: str = "TMC_NA_template_FR.docx"):
        """GÃ©nÃ©rer le CV TMC final avec docxtpl"""
        print(f"ğŸ“ GÃ©nÃ©ration du CV TMC: {output_path}")
        
        # ğŸ” RECHERCHE INTELLIGENTE DU TEMPLATE (nouvelle fonction robuste)
        final_template_path = self.find_template_file(template_path)
        print(f"   ğŸ“„ Template: {final_template_path}")
        
        # CrÃ©er environnement Jinja2 avec filtre pairwise
        jinja_env = jinja2.Environment()
        
        def pairwise(iterable):
            items = list(iterable)
            result = []
            for i in range(0, len(items), 2):
                if i + 1 < len(items):
                    result.append((items[i], items[i + 1]))
                else:
                    result.append((items[i], ''))
            return result
        
        jinja_env.filters['pairwise'] = pairwise
        
        # ğŸ”¥ Ajouter la fonction r pour RichText dans le contexte
        context['r'] = lambda x: x
        
        # âš ï¸ CORRECTION XML : Ã‰chapper les caractÃ¨res spÃ©ciaux (Â®, &, <, >, etc.)
        from html import escape as html_escape
        print(f"   ğŸ”§ Ã‰chappement des caractÃ¨res XML spÃ©ciaux...")
        
        # Ã‰chapper les champs texte simples
        for key in ['first_name', 'last_name', 'title', 'FIRST_NAME', 'LAST_NAME', 
                   'TITLE', 'residency', 'RESIDENCY', 'languages', 'LANGUAGES']:
            if key in context and isinstance(context[key], str):
                context[key] = html_escape(context[key])
        
        # Ã‰chapper les expÃ©riences
        if 'work_experience' in context:
            for exp in context['work_experience']:
                for key in ['period', 'company', 'position']:
                    if key in exp and isinstance(exp[key], str):
                        exp[key] = html_escape(exp[key])
                
                if 'general_responsibilities' in exp and isinstance(exp['general_responsibilities'], list):
                    exp['general_responsibilities'] = [
                        html_escape(r) if isinstance(r, str) else r
                        for r in exp['general_responsibilities']
                    ]
        
        # Ã‰chapper formations
        if 'education' in context:
            for edu in context['education']:
                for key in ['institution', 'degree', 'graduation_year', 'country', 'level', 'title']:
                    if key in edu and isinstance(edu[key], str):
                        edu[key] = html_escape(edu[key])
        
        # Ã‰chapper certifications
        if 'certifications' in context:
            for cert in context['certifications']:
                for key in ['name', 'institution', 'year', 'country']:
                    if key in cert and isinstance(cert[key], str):
                        cert[key] = html_escape(cert[key])
        
        # Ã‰chapper projets
        if 'projects' in context:
            for proj in context['projects']:
                for key in ['nom', 'description']:
                    if key in proj and isinstance(proj[key], str):
                        proj[key] = html_escape(proj[key])
        
        print(f"   âœ… CaractÃ¨res XML Ã©chappÃ©s (Â®, &, <, >, etc.)")
        
        # Charger le template TMC
        doc = DocxTemplate(final_template_path)
        
        # Rendre le document
        doc.render(context, jinja_env)
        
        # Sauvegarder
        doc.save(output_path)
        print(f"âœ… CV TMC gÃ©nÃ©rÃ© avec succÃ¨s!")

    def generate_ms_cv_3parts(self, tmc_context, skills_matrix_path, output_path, 
                              cover_template="TMC_NA_template_EN_Anonymise_CoverPage.docx",
                              content_template="TMC_NA_template_EN_Anonymise_Content.docx"):
        """
        GÃ©nÃ¨re un CV Morgan Stanley en 3 parties:
        1. Cover page (photo + nom + titre + location + langues)
        2. Skills Matrix (uploadÃ©e par le recruteur)
        3. Contenu dÃ©taillÃ© (profile + skills + experiences + education)
        
        Args:
            tmc_context: Contexte enrichi du candidat
            skills_matrix_path: Path vers le fichier Skills Matrix uploadÃ©
            output_path: Path pour le fichier final
            cover_template: Template pour la cover page
            content_template: Template pour le contenu dÃ©taillÃ©
        
        Returns:
            tuple: (success: bool, output_path: str)
        """
        try:
            from pathlib import Path
            from docxcompose.composer import Composer
            from docx import Document
            import shutil
            
            # Dossier temporaire
            temp_dir = Path("/tmp/cv_optimizer_ms")
            temp_dir.mkdir(exist_ok=True)
            
            # Ã‰TAPE 1: GÃ©nÃ©rer cover page
            print("ğŸ¨ Generating cover page...")
            cover_path = temp_dir / "cover.docx"
            
            # âœ… FIX: Passer seulement le nom du template, find_template_file va le chercher
            print(f"   ğŸ“„ Using cover template: {cover_template}")
            
            self.generate_tmc_docx(
                tmc_context, 
                str(cover_path), 
                template_path=cover_template  # Juste le nom, pas le chemin complet
            )
            print(f"   âœ… Cover page generated: {cover_path.name}")
            
            # Ã‰TAPE 2: Merger cover + Skills Matrix
            print("ğŸ”— Merging cover with Skills Matrix...")
            cover_with_skills = temp_dir / "cover_and_skills.docx"
            
            # Charger les deux documents
            cover_doc = Document(str(cover_path))
            skills_doc = Document(skills_matrix_path)
            
            # âœ… V1.3.4.2 FIX: Change table width from fixed to auto to prevent horizontal shift
            print("ğŸ”§ Fixing Skills Matrix table width...")
            tables_fixed = fix_table_width_to_auto(skills_doc)
            print(f"   âœ… Fixed {tables_fixed} table(s) to auto width")
            
            # V1.3.4 FIX: Ajuster les marges de la Skills Matrix pour correspondre au template
            # Copier les marges du cover vers skills avant merge
            cover_sections = cover_doc.sections
            skills_sections = skills_doc.sections
            
            if cover_sections and skills_sections:
                # Utiliser les marges du template pour la Skills Matrix
                for section in skills_sections:
                    section.top_margin = cover_sections[0].top_margin
                    section.bottom_margin = cover_sections[0].bottom_margin
                    section.left_margin = cover_sections[0].left_margin
                    section.right_margin = cover_sections[0].right_margin
            
            # V1.3.4.1 FIX: Supprimer les espacements au dÃ©but de la Skills Matrix
            # Ceci assure que le contenu commence exactement en haut de la page
            from docx.shared import Pt
            from docx.oxml import parse_xml
            
            # Supprimer TOUS les paragraphes vides au dÃ©but du body XML
            # Travailler directement sur body._element pour avoir l'ordre exact
            body = skills_doc.element.body
            elements_to_remove = []
            
            # Parcourir les Ã©lÃ©ments dans l'ordre et marquer les paragraphes vides au dÃ©but
            for elem in body:
                tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
                
                if tag == 'p':  # C'est un paragraphe
                    # VÃ©rifier s'il est vide (pas de texte)
                    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                    text_elems = elem.findall('.//w:t', ns)
                    text_content = ''.join([t.text for t in text_elems if t.text])
                    
                    if not text_content.strip():
                        # Paragraphe vide au dÃ©but â†’ marquer pour suppression
                        elements_to_remove.append(elem)
                    else:
                        # Premier paragraphe avec texte â†’ arrÃªter
                        break
                elif tag == 'tbl':
                    # On a atteint une table â†’ arrÃªter
                    break
            
            # Supprimer les Ã©lÃ©ments marquÃ©s
            for elem in elements_to_remove:
                body.remove(elem)
            
            print(f"   ğŸ§¹ Removed {len(elements_to_remove)} empty paragraphs from Skills Matrix")
            
            # RÃ©initialiser le spacing du premier Ã©lÃ©ment restant (si paragraphe)
            if skills_doc.paragraphs:
                first_para = skills_doc.paragraphs[0]
                first_para.paragraph_format.space_before = Pt(0)
                first_para.paragraph_format.space_after = Pt(0)
            
            # Ajouter page break aprÃ¨s cover
            cover_doc.add_page_break()
            
            # Merger avec docxcompose
            composer = Composer(cover_doc)
            composer.append(skills_doc)
            
            # Sauvegarder
            composer.save(str(cover_with_skills))
            print(f"   âœ… Cover + Skills Matrix merged")
            
            # Ã‰TAPE 3: GÃ©nÃ©rer contenu dÃ©taillÃ©
            print("ğŸ“ Generating detailed content...")
            content_path = temp_dir / "content.docx"
            
            # âœ… FIX: Passer seulement le nom du template
            print(f"   ğŸ“„ Using content template: {content_template}")
            
            self.generate_tmc_docx(
                tmc_context,
                str(content_path),
                template_path=content_template  # Juste le nom, pas le chemin complet
            )
            print(f"   âœ… Content generated: {content_path.name}")
            
            # Ã‰TAPE 4: Merger tout ensemble
            print("ğŸ”— Merging everything...")
            
            # Charger cover+skills
            final_doc = Document(str(cover_with_skills))
            
            # Ajouter page break avant content
            final_doc.add_page_break()
            
            # Merger avec content
            final_composer = Composer(final_doc)
            content_doc = Document(str(content_path))
            final_composer.append(content_doc)
            
            # Sauvegarder le document final
            final_composer.save(str(output_path))
            print(f"âœ… Final CV saved: {output_path}")
            
            # Nettoyer les fichiers temporaires
            shutil.rmtree(temp_dir, ignore_errors=True)
            
            return True, str(output_path)
            
        except Exception as e:
            error_msg = f"Error generating MS CV: {str(e)}"
            print(f"âŒ {error_msg}")
            import traceback
            traceback.print_exc()
            return False, error_msg
    def apply_bold_post_processing(self, docx_path: str, keywords: list):
        """Post-traiter le document pour mettre en gras les technologies dans les tableaux"""
        print(f"ğŸ¨ Application du gras sur les technologies...")
        
        from docx import Document as DocxDocument
        from docx.shared import RGBColor
        import re
        
        doc = DocxDocument(docx_path)
        modifications = 0
        
        print(f"   Recherche des **mot** dans le document...")
        
        def apply_bold_to_runs(paragraph):
            """Trouve **mot** et met en gras UNIQUEMENT ce mot"""
            text = paragraph.text
            if '**' not in text:
                return 0
            
            changes = 0
            # Pattern pour trouver **mot**
            pattern = re.compile(r'\*\*([^*]+)\*\*')
            
            # Reconstituer le paragraphe avec le bon formatage
            matches = list(pattern.finditer(text))
            if not matches:
                return 0
            
            # Supprimer tous les runs existants
            for run in paragraph.runs:
                run._element.getparent().remove(run._element)
            
            # Reconstruire avec le bon formatage
            last_end = 0
            for match in matches:
                # Texte normal avant
                if match.start() > last_end:
                    run = paragraph.add_run(text[last_end:match.start()])
                    run.bold = False
                    run.font.name = 'Arial'
                
                # Texte en gras
                run = paragraph.add_run(match.group(1))
                run.bold = True
                run.font.name = 'Arial'
                changes += 1
                
                last_end = match.end()
            
            # Texte normal aprÃ¨s
            if last_end < len(text):
                run = paragraph.add_run(text[last_end:])
                run.bold = False
                run.font.name = 'Arial'
            
            return changes
        
        # Parcourir TOUS les tableaux (oÃ¹ sont les expÃ©riences)
        print("   ğŸ“‹ Traitement des tableaux...")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        modifications += apply_bold_to_runs(paragraph)
        
        # Parcourir aussi les paragraphes normaux
        print("   ğŸ“ Traitement des paragraphes...")
        for paragraph in doc.paragraphs:
            modifications += apply_bold_to_runs(paragraph)
        
        # Sauvegarder
        doc.save(docx_path)
        if modifications > 0:
            print(f"âœ… {modifications} mots mis en gras")
        else:
            print(f"âš ï¸ Aucun **mot** trouvÃ©")
        
        return modifications
        
def main():
    """Point d'entrÃ©e CLI"""
    import argparse
    
    parser = argparse.ArgumentParser(description='TMC Universal CV Enricher')
    parser.add_argument('cv_path', help='Chemin du CV (PDF, Word, etc.)')
    parser.add_argument('jd_path', help='Chemin de la Job Description')
    parser.add_argument('--output', '-o', default='cv_enriched_tmc.docx', help='Fichier de sortie')
    
    args = parser.parse_args()
    
    try:
        enricher = TMCUniversalEnricher()
        
        print("\nğŸš€ TMC UNIVERSAL CV ENRICHER")
        print("=" * 60)
        
        # MODULE 1: Extraction
        print("\n[1/5] Extraction du CV...")
        cv_text = enricher.extract_cv_text(args.cv_path)
        print(f"      âœ… {len(cv_text)} caractÃ¨res extraits")
        
        # MODULE 2: Parsing
        print("\n[2/5] Parsing intelligent...")
        parsed_cv = enricher.parse_cv_with_claude(cv_text)
        
        # MODULE 3: Enrichissement
        print("\n[3/5] Enrichissement avec IA...")
        jd_text = enricher.read_job_description(args.jd_path)
        enriched_cv = enricher.enrich_cv_with_prompt(parsed_cv, jd_text)
        
        # MODULE 4: Mapping TMC
        print("\n[4/5] Mapping structure TMC...")
        tmc_context = enricher.map_to_tmc_structure(parsed_cv, enriched_cv)
        
        # MODULE 5: GÃ©nÃ©ration
        print("\n[5/5] GÃ©nÃ©ration CV final...")
        enricher.generate_tmc_docx(tmc_context, args.output)
        
        # POST-PROCESSING: Application du gras
        print("\n[POST] Application du gras sur mots-clÃ©s...")
        keywords = enriched_cv.get('mots_cles_a_mettre_en_gras', [])
        print(f"   Mots-clÃ©s Ã  mettre en gras: {keywords}")
        
        if keywords:
            result = enricher.apply_bold_post_processing(args.output, keywords)
            if result == 0:
                print("   âš ï¸ AUCUN mot-clÃ© n'a Ã©tÃ© mis en gras!")
                print("   VÃ©rifiez que les mots-clÃ©s sont bien dans le CV")
        else:
            print("   âš ï¸ Aucun mot-clÃ© retournÃ© par l'IA")
        
        # RÃ‰SUMÃ‰ FINAL
        print("\n" + "=" * 60)
        print("ğŸ‰ ENRICHISSEMENT TERMINÃ‰!")
        print("=" * 60)
        print(f"ğŸ“Š Score matching: {enriched_cv.get('score_matching', 0)}/100")
        
        # Afficher les domaines analysÃ©s
        if enriched_cv.get('domaines_analyses'):
            print(f"\nğŸ“Š Analyse par domaine:")
            for domaine in enriched_cv['domaines_analyses']:
                match = domaine.get('match', '')
                emoji = 'âŒ' if match == 'incompatible' else 'âš ï¸' if match == 'partiel' else 'âœ…'
                print(f"   {emoji} {domaine.get('domaine', 'N/A')}: {domaine.get('score', 0)}/{domaine.get('score_max', 0)} pts ({domaine.get('poids', 0)}%)")
                print(f"      â†’ {domaine.get('commentaire', 'N/A')}")
        
        if enriched_cv.get('synthese_matching'):
            print(f"\nğŸ’¬ SynthÃ¨se: {enriched_cv['synthese_matching']}")
        
        print(f"\nğŸ’ª Points forts:")
        for pf in enriched_cv.get('points_forts', [])[:3]:
            print(f"   â€¢ {pf}")
        print(f"\nğŸ“„ Fichier gÃ©nÃ©rÃ©: {args.output}")
        print("=" * 60)
        
    except Exception as e:
        print(f"\nâŒ ERREUR: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
