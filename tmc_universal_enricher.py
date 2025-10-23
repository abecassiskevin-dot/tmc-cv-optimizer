#!/usr/bin/env python3
"""
TMC Universal CV Enricher
Lit n'importe quel CV ‚Üí Enrichit avec IA ‚Üí G√©n√®re CV TMC professionnel
"""

import os
import sys
import json
from docxtpl import DocxTemplate, RichText
from docx import Document
import jinja2
from typing import Dict, List, Any
import PyPDF2
import re
from zipfile import ZipFile
from xml.etree import ElementTree as ET

print(">>> tmc_universal_enricher module loading", flush=True)


class TMCUniversalEnricher:
    """Enrichisseur universel de CV au format TMC"""
    
    def __init__(self, api_key: str = None):
        """Initialiser avec cl√© API Claude"""
        self.api_key = api_key or os.getenv('ANTHROPIC_API_KEY')
        if not self.api_key:
            raise ValueError("‚ùå Cl√© API Claude manquante! D√©finissez ANTHROPIC_API_KEY dans les secrets Streamlit ou en variable d'environnement.")
        
        # Debug cl√© API
        print(f">>> ANTHROPIC_KEY_PRESENT: {bool(self.api_key)}, len: {len(self.api_key) if self.api_key else 0}", flush=True)
        
        # Ne cr√©e PAS le client ici (lazy loading)
        self._anthropic_client = None
    
    def _get_anthropic_client(self):
        """Lazy loading du client Anthropic"""
        if self._anthropic_client is None:
            try:
                print(">>> Creating anthropic client", flush=True)
                import anthropic
                # Cr√©ation SIMPLE du client pour version 0.25.9
                self._anthropic_client = anthropic.Anthropic(api_key=self.api_key)
                print(">>> Anthropic client created OK", flush=True)
            except Exception as e:
                print(f">>> ERROR creating anthropic client: {repr(e)}", flush=True)
                raise
        return self._anthropic_client
    
    # ========================================
    # MODULE 1 : EXTRACTION UNIVERSELLE
    # ========================================
    
    def detect_file_type(self, file_path: str) -> str:
        """D√©tecter le type de fichier"""
        ext = file_path.lower().split('.')[-1]
        if ext == 'pdf':
            return 'pdf'
        elif ext in ['docx', 'doc']:
            return 'docx'
        elif ext in ['txt', 'text']:
            return 'txt'
        else:
            return 'unknown'
    
    def extract_from_pdf(self, file_path: str) -> str:
        """Extraire texte d'un PDF"""
        try:
            text = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
            return "\n".join(text)
        except Exception as e:
            print(f"‚ö†Ô∏è Erreur extraction PDF: {e}")
            return ""
    
     
    def extract_from_docx(self, file_path: str) -> str:
        """Extraire texte d'un Word + zones textes"""
        try:
            doc = Document(file_path)
            text = []
            
            # Paragraphes normaux
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text.strip())
            
            # Tableaux
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text.append(row_text)
            
            # NOUVEAU : Extraire les zones textes du XML
            textbox_content = self.extract_textboxes(file_path)
            if textbox_content:
                text.append("\n=== ZONES TEXTES ===")
                text.extend(textbox_content)
            
            return "\n".join(text)
        except Exception as e:
            print(f"‚ö†Ô∏è Erreur extraction Word: {e}")
            return ""
    def extract_from_txt(self, file_path: str) -> str:
        """Extraire texte d'un fichier texte"""
        try:
            # Essayer plusieurs encodages
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            # Si tout √©choue, ignorer les erreurs
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            print(f"‚ö†Ô∏è Erreur extraction TXT: {e}")
            return ""
    
    def extract_textboxes(self, docx_path: str) -> list:
        """Extraire le contenu des zones textes (text boxes) du XML"""
        textboxes = []
        
        try:
            with ZipFile(docx_path, 'r') as docx:
                # Lire le document.xml
                xml_content = docx.read('word/document.xml')
                tree = ET.fromstring(xml_content)
                
                # Namespaces Word
                namespaces = {
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                    'v': 'urn:schemas-microsoft-com:vml',
                    'w10': 'urn:schemas-microsoft-com:office:word'
                }
                
                # Chercher tous les √©l√©ments de texte dans les zones textes
                # Les zones textes sont dans w:txbxContent
                for txbx in tree.findall('.//w:txbxContent', namespaces):
                    texts = []
                    for t in txbx.findall('.//w:t', namespaces):
                        if t.text:
                            texts.append(t.text.strip())
                    if texts:
                        textboxes.append(' '.join(texts))
                
                # Aussi chercher dans v:textbox (ancien format)
                for vtxbx in tree.findall('.//v:textbox', namespaces):
                    texts = []
                    for t in vtxbx.findall('.//w:t', namespaces):
                        if t.text:
                            texts.append(t.text.strip())
                    if texts:
                        textboxes.append(' '.join(texts))
                        
        except Exception as e:
            print(f"‚ö†Ô∏è Erreur extraction zones textes: {e}")
        
        return textboxes
    
    def extract_cv_text(self, cv_path: str) -> str:
        """Extraction universelle - d√©tecte et extrait selon le type"""
        print(f"üìÑ Extraction du CV: {cv_path}")
        
        file_type = self.detect_file_type(cv_path)
        
        if file_type == 'pdf':
            print("   Format d√©tect√©: PDF")
            return self.extract_from_pdf(cv_path)
        elif file_type == 'docx':
            print("   Format d√©tect√©: Word")
            return self.extract_from_docx(cv_path)
        elif file_type == 'txt':
            print("   Format d√©tect√©: Texte")
            return self.extract_from_txt(cv_path)
        else:
            raise ValueError(f"‚ùå Format non support√©: {file_type}")

    # ========================================
    # MODULE 2 : PARSING INTELLIGENT
    # ========================================
    
    def parse_cv_with_claude(self, cv_text: str) -> Dict[str, Any]:
        """Parser le CV avec Claude pour extraire les infos structur√©es"""
        print("ü§ñ Parsing du CV avec Claude AI...", flush=True)
        
        try:
            client = self._get_anthropic_client()
            
            prompt = f"""Tu es un expert en analyse de CV. Extrait TOUTES les informations de ce CV et structure-les en JSON.

CV √Ä ANALYSER:
{cv_text}

IMPORTANT CRITIQUE:
- Le NOM peut √™tre cach√© dans un tableau HTML ou √™tre stylis√©. Cherche PARTOUT.
- Le LIEU DE R√âSIDENCE est OBLIGATOIRE : cherche "Montr√©al", "Montreal", villes + pays (ex: "Montreal CA", "Montr√©al, Canada", "Toronto ON", etc.). Si introuvable, mets "Location not specified".
- Les LANGUES sont OBLIGATOIRES : cherche "Fran√ßais", "French", "English", "Anglais", "Bilingual", "Bilingue", etc. Si introuvable, mets ["Not specified"].

Extrait et structure en JSON STRICT (sans markdown):
{{
  "nom_complet": "Nom Pr√©nom du candidat (cherche PARTOUT, m√™me dans tableaux/HTML)",
  "titre_professionnel": "Titre/poste actuel",
  "profil_resume": "R√©sum√© du profil si pr√©sent (sinon vide)",
  "lieu_residence": "OBLIGATOIRE - Ville, Pays (ex: Montr√©al, Canada) ou Montreal CA. Cherche codes pays (CA, US, FR). Si vraiment introuvable: 'Location not specified'",
  "langues": ["OBLIGATOIRE - Fran√ßais", "Anglais", ... Cherche 'bilingual', 'French', 'English', etc. Si introuvable: ['Not specified']],
  "competences": ["comp√©tence1", "comp√©tence2", "comp√©tence3", ...],
  "experiences": [
    {{
      "periode": "2020-2023",
      "entreprise": "Nom entreprise",
      "poste": "Titre du poste",
      "responsabilites": ["t√¢che 1", "t√¢che 2", "t√¢che 3"]
    }}
  ],
  "formation": [
    {{
      "diplome": "Nom COMPLET du dipl√¥me",
      "institution": "Nom √©cole/universit√©",
      "annee": "2020 (ou p√©riode exacte)",
      "pays": "Canada"
    }}
  ],
  "certifications": [
    {{
      "nom": "Nom certification",
      "organisme": "Organisme",
      "annee": "2023"
    }}
  ],
  "projets": [
    {{
      "nom": "Nom projet",
      "description": "Description courte"
    }}
  ]
}}

R√àGLES CRITIQUES:
- Le NOM est PRIORITAIRE - cherche dans tout le texte (tableaux, d√©but, fin)
- LIEU DE R√âSIDENCE : cherche formats "Ville, Pays", "Montreal CA", "Montr√©al QC", codes postaux (H2X, etc.)
- LANGUES : cherche "Languages", "Langues", "French", "English", "Bilingual", m√™me dans sections comp√©tences
- Pour les dipl√¥mes: nom COMPLET + ann√©e EXACTE
- Extrait TOUT (ne rate rien)
- Si une section est vide, mets une liste vide []
- Format JSON strict uniquement"""

            response = client.messages.create(
                model="claude-sonnet-4-5-20250929",
                max_tokens=8000,
                messages=[{"role": "user", "content": prompt}]
            )
            
        except Exception as e:
            print(f">>> ERROR calling anthropic for parsing: {repr(e)}", flush=True)
            return {}
        
        response_text = response.content[0].text.strip()
        
        # Nettoyer JSON
        if response_text.startswith('```json'):
            response_text = response_text[7:]
        if response_text.startswith('```'):
            response_text = response_text[3:]
        if response_text.endswith('```'):
            response_text = response_text[:-3]
        response_text = response_text.strip()
        
        try:
            parsed_data = json.loads(response_text)
            print(f"‚úÖ Parsing r√©ussi!")
            print(f"   Nom: [ANONYMIZED]")
            print(f"   Langues: {', '.join(parsed_data.get('langues', []))}")
            print(f"   Lieu: [ANONYMIZED]")
            print(f"   Comp√©tences: {len(parsed_data.get('competences', []))}")
            print(f"   Exp√©riences: {len(parsed_data.get('experiences', []))}")
            return parsed_data
        except json.JSONDecodeError as e:
            print(f"‚ö†Ô∏è Erreur JSON: {e}")
            print(f"R√©ponse brute: {response_text[:500]}")
            return {}

    # ========================================
    # MODULE 3 : ENRICHISSEMENT (TON PROMPT)
    # ========================================
    
    def read_job_description(self, jd_path: str) -> str:
        """Lire la job description"""
        file_type = self.detect_file_type(jd_path)
        
        if file_type == 'pdf':
            return self.extract_from_pdf(jd_path)
        elif file_type == 'docx':
            return self.extract_from_docx(jd_path)
        else:
            return self.extract_from_txt(jd_path)
    
    def enrich_cv_with_prompt(self, parsed_cv: Dict[str, Any], jd_text: str, language: str = "French") -> Dict[str, Any]:
        """Enrichir le CV avec ton prompt exact"""
        print(f"‚ú® Enrichissement du CV avec l'IA...", flush=True)
        print(f"   Langue cible: {language}", flush=True)
        
        try:
            client = self._get_anthropic_client()
            
            # Reconstruire le CV en texte pour le prompt
            cv_text = f"""
PROFIL: {parsed_cv.get('profil_resume', '')}

TITRE: {parsed_cv.get('titre_professionnel', '')}

COMP√âTENCES:
{chr(10).join(['- ' + comp for comp in parsed_cv.get('competences', [])])}

EXP√âRIENCES:
"""
            for exp in parsed_cv.get('experiences', []):
                cv_text += f"\n{exp.get('periode', '')} | {exp.get('entreprise', '')} | {exp.get('poste', '')}\n"
                for resp in exp.get('responsabilites', []):
                    cv_text += f"  - {resp}\n"
            
            cv_text += "\nFORMATION:\n"
            for form in parsed_cv.get('formation', []):
                cv_text += f"- {form.get('diplome', '')} | {form.get('institution', '')} | {form.get('annee', '')}\n"
        
            # PROMPT ULTRA-RENFORC√â POUR COH√âRENCE ABSOLUE
            language_instruction = f"""
‚ö†Ô∏è R√àGLE ABSOLUE - LANGUE {language.upper()}:
- Tu DOIS g√©n√©rer 100% du contenu en {language}
- Le TITRE PROFESSIONNEL doit √™tre en {language}
- TOUTES les descriptions doivent √™tre en {language}
- TOUS les mots-cl√©s doivent √™tre en {language}
- Respecte les conventions professionnelles de la langue {language}
- Si {language} = French: utilise "Analyste", "Gestion", "Configuration", etc.
- Si {language} = English: utilise "Analyst", "Management", "Configuration", etc.

IMPORTANT TITRE:
- Adapte le titre professionnel √† la Job Description
- Le titre doit √™tre COURT (3-5 mots maximum)
- Le titre doit √™tre en {language}
- Exemple en fran√ßais: "Analyste QA Senior" ou "Analyste Configuration SharePoint"
- Exemple en anglais: "Senior QA Analyst" or "SharePoint Configuration Analyst"

üéØ R√îLE CRITIQUE - TU ES UN RECRUTEUR SENIOR PROFESSIONNEL:
- Tu as 15+ ans d'exp√©rience en recrutement technique
- Tu travailles pour le CLIENT (l'entreprise qui recrute)
- Ta mission: √©valuer si le CANDIDAT correspond EXACTEMENT aux besoins du CLIENT
- Tu dois √™tre OBJECTIF, RIGOUREUX et REPRODUCTIBLE dans ton √©valuation
- Ton scoring doit √™tre IDENTIQUE si tu analyses le m√™me CV/JD plusieurs fois
- Tu notes comme un examinateur professionnel, pas comme un vendeur
"""
            
            prompt = f"""Voici la job description et le CV actuel ci-dessous.

üîπ Am√©liore le CV pour qu'il soit parfaitement align√© avec la job description tout en gardant le format d'origine (titres, mise en page, structure, ton professionnel).
{language_instruction}

üéØ ANALYSE DE MATCHING POND√âR√âE (ULTRA-CRITIQUE - COH√âRENCE ABSOLUE REQUISE):

‚ö†Ô∏è PRINCIPE FONDAMENTAL DE COH√âRENCE - M√âTHODOLOGIE STRICTE:
- Tu es un SYST√àME D'√âVALUATION AUTOMATIS√â, pas un humain
- Pour le M√äME CV et la M√äME JD ‚Üí EXACTEMENT le m√™me score √† chaque fois
- Utilise une grille d'√©valuation MATH√âMATIQUE et REPRODUCTIBLE
- Agis comme un ALGORITHME, pas comme un recruteur subjectif
- Chaque crit√®re suit des r√®gles BINAIRES strictes (oui/non, pr√©sent/absent)
- Tu DOIS pouvoir justifier CHAQUE point attribu√© avec des FAITS du CV
- Si tu h√©sites entre 2 scores ‚Üí prends le PLUS BAS (principe de strictness)

üî¥ R√àGLE D'OR - SCORE GLOBAL = SOMME DOMAINES:
- Le score_matching FINAL = somme EXACTE de tous les scores de domaines
- V√âRIFIE 3 FOIS avant de r√©pondre: somme des scores = score_matching
- Si tu calcules 37/100 en sommant les domaines ‚Üí score_matching DOIT √™tre 37
- NE JAMAIS inventer un score global diff√©rent de la somme calcul√©e

√âTAPE 1 - IDENTIFIER 5-8 DOMAINES CRITIQUES (M√âTHODE ALGORITHIMQUE):

üìã PROCESSUS AUTOMATIQUE D'IDENTIFICATION:
1. Scan complet de la JD - rep√©rer TOUS les mots techniques
2. Compter la fr√©quence EXACTE de chaque technologie/comp√©tence
3. Cr√©er une liste de domaines par ordre d'importance
4. Appliquer la formule de pond√©ration ci-dessous

üìä FORMULE DE POND√âRATION MATH√âMATIQUE:
Pour chaque domaine, calcule son poids avec:

Poids_Base = (Nombre_mentions / Total_mentions_techniques) √ó 100

Bonus:
- +20% si c'est le TITRE du poste (ex: ".NET Developer" ‚Üí Stack .NET = +20%)
- +15% si mots "Required", "Must have", "Essential", "Critical"
- +10% si mentionn√© dans les 3 premi√®res lignes de la JD
- +5% par occurrence au-del√† de 3 mentions

Poids_Final = min(Poids_Base + Bonus, 50%)  ‚Üê Aucun domaine ne peut d√©passer 50%

R√àGLES STRICTES DE POND√âRATION:
- Stack technique principal (dans titre ou 5+ mentions): 30-50%
- Comp√©tences techniques secondaires (3-4 mentions): 15-25%
- Comp√©tences techniques tertiaires (1-2 mentions): 5-15%
- Soft skills/Leadership: MAXIMUM 10% (sauf si poste management)
- TOTAL des poids = EXACTEMENT 100% (v√©rifie avec calculatrice)
- Si total ‚â† 100%, ajuste proportionnellement tous les poids

√âTAPE 2 - SCORER CHAQUE DOMAINE (ALGORITHME DE NOTATION STRICT):

ü§ñ SYST√àME DE NOTATION AUTOMATIS√â - APPLIQUE CES R√àGLES EXACTEMENT:

POUR CHAQUE DOMAINE, COMPTE:
a) Nombre de mentions EXACTES de la technologie dans le CV
b) Nombre de projets/exp√©riences utilisant cette technologie  
c) Dur√©e totale d'utilisation (ann√©es)
d) Niveau d√©montr√© (junior/interm√©diaire/senior)

üìê FORMULE MATH√âMATIQUE DE SCORING:

√âtape 2A - Score Brut (0-100%):
‚Ä¢ 0% : Z√âRO mention de la techno dans le CV, stack incompatible
‚Ä¢ 10% : Technologie proche mentionn√©e (PostgreSQL pour SQL Server)
‚Ä¢ 25% : 1 mention + aucune exp√©rience pratique (formation seulement)
‚Ä¢ 40% : 1-2 mentions + 1 projet + <1 an d'exp√©rience
‚Ä¢ 60% : 3-4 mentions + 2 projets + 1-2 ans d'exp√©rience
‚Ä¢ 80% : 5+ mentions + 3+ projets + 3+ ans d'exp√©rience
‚Ä¢ 100% : 7+ mentions + expertise d√©montr√©e + senior confirm√©

√âtape 2B - Ajustements OBLIGATOIRES:
‚Ä¢ Si stack incompatible (Java vs .NET) ‚Üí Score = 0% (NON-N√âGOCIABLE)
‚Ä¢ Si technologie absente du CV ‚Üí Score = 0% (NON-N√âGOCIABLE)
‚Ä¢ Si aucune exp√©rience pratique prouv√©e ‚Üí Score MAX = 30%
‚Ä¢ Si exp√©rience <1 an ‚Üí Score MAX = 50%
‚Ä¢ Si niveau junior √©vident ‚Üí Score MAX = 60%

√âtape 2C - Calcul Final:
Score_Domaine = (Score_Brut √ó Poids_Domaine) / 100

EXEMPLE D√âTAILL√â:
Domaine: ".NET Development" - Poids: 40%
CV candidat: AUCUNE mention .NET, seulement Java
‚Üí Score_Brut = 0%
‚Üí Score_Domaine = (0 √ó 40) / 100 = 0 points
‚Üí Commentaire: "‚ùå Stack incompatible - profil Java exclusif"

üî¥ V√âRIFICATION FINALE OBLIGATOIRE:
Somme_Scores = Œ£(tous les Score_Domaine)
Si Somme_Scores ‚â† score_matching ‚Üí ERREUR CRITIQUE ‚Üí RECALCULE

√âTAPE 3 - COMMENTAIRE PAR DOMAINE (30-50 mots):
- Utilise ‚ùå (0-30%), ‚ö†Ô∏è (30-70%), ‚úÖ (70-100%)
- Sois FACTUEL et OBJECTIF dans tes commentaires
- Base-toi UNIQUEMENT sur les FAITS pr√©sents dans le CV
- Ne fais PAS d'hypoth√®ses optimistes

EXEMPLE:
JD demande: ".NET, C#, Azure, SQL Server"
Candidat a: "Java, AWS, PostgreSQL"

R√âSULTAT:
{{
  "domaines_analyses": [
    {{
      "domaine": "Stack .NET (C#, ASP.NET Core, Entity Framework)",
      "poids": 40,
      "score": 0,
      "score_max": 40,
      "commentaire": "‚ùå Aucune exp√©rience .NET/C#. Profil Java exclusivement - incompatibilit√© majeure sur stack principale.",
      "match": "incompatible"
    }},
    {{
      "domaine": "Cloud Microsoft Azure",
      "poids": 20,
      "score": 8,
      "score_max": 20,
      "commentaire": "‚ö†Ô∏è Exp√©rience AWS uniquement. Comp√©tences cloud transf√©rables mais n√©cessite formation Azure.",
      "match": "partiel"
    }},
    {{
      "domaine": "SQL Server & T-SQL",
      "poids": 15,
      "score": 10,
      "score_max": 15,
      "commentaire": "‚úÖ Ma√Ætrise PostgreSQL et MySQL - comp√©tences SQL transf√©rables √† SQL Server.",
      "match": "bon"
    }}
  ],
  "score_matching": 45,
  "synthese_matching": "Profil Java senior inadapt√© pour poste .NET. Gap critique sur stack principale (0/40). Comp√©tences transf√©rables en cloud et SQL, mais n√©cessite reconversion majeure."
}}

Fais :

1. ANALYSE POND√âR√âE OBLIGATOIRE (voir ci-dessus)
2. Une version r√©√©crite et enrichie du CV

2b. PROFIL exceptionnel : √©cris un paragraphe NARRATIF fluide (pas de liste), 5-6 lignes avec progression logique.

2c. GRAS ULTRA-S√âLECTIF : identifie UNIQUEMENT 3-5 technologies CRITIQUES.

3. Int√®gre naturellement les mots-cl√©s techniques de la JD
4. Ajuste les intitul√©s pour que le profil paraisse livrable imm√©diatement
5. N'invente rien ‚Äî reformule uniquement les √©l√©ments pr√©sents
6. EXP√âRIENCES : bullets courts (1 ligne max), maximum 5-6 bullets par exp√©rience

R√©ponds en JSON STRICT (sans markdown) avec cette structure:
{{
  "domaines_analyses": [
    {{
      "domaine": "Nom domaine technique/fonctionnel (ex: Stack .NET, Cloud Azure)",
      "poids": 40,
      "score": 15,
      "score_max": 40,
      "commentaire": "Explication 30-50 mots avec ‚ùå/‚ö†Ô∏è/‚úÖ",
      "match": "incompatible|partiel|bon|excellent"
    }}
  ],
  "score_matching": 45,
  "synthese_matching": "R√©sum√© 2-3 phrases du matching global avec points forts et gaps critiques",
  
  "titre_professionnel_enrichi": "TITRE COURT en {language} (3-5 mots max)",
  
  "profil_enrichi": "Profil NARRATIF 5-6 lignes en {language} avec **3-5 technologies cl√©s** en gras",
  
  "mots_cles_a_mettre_en_gras": ["Liste 15-20 TECHNOLOGIES de la JD - PAS de verbes g√©n√©riques"],
  
  "competences_enrichies": {{
    "Nom Cat√©gorie 1 (3-6 mots max)": [
      "**Technologie principale** : description en 2-3 lignes (MAXIMUM 100-150 caract√®res) incluant contexte, outils associ√©s (**outil1**, **outil2**) et r√©sultats. Style concis et percutant.",
      "**Autre technologie** : description COURTE avec contexte + outils (**tech1**, **tech2**) + impact. 2-3 technologies en **gras** par comp√©tence."
    ],
    "Nom Cat√©gorie 2": [
      "Comp√©tence concise..."
    ]
  }},
  
  R√àGLES ULTRA-CRITIQUES pour les comp√©tences (NON-N√âGOCIABLE):
  - Noms de cat√©gories COURTS (3-6 mots max)
  - 5-6 cat√©gories ADAPT√âES √† la JD
  - Chaque cat√©gorie: 3-5 comp√©tences MAXIMUM
  - CHAQUE comp√©tence : 2-3 LIGNES MAXIMUM (100-150 caract√®res) - NE PAS D√âPASSER
  - Format: "**Technologie** : description concise avec outils (**outil1**, **outil2**) + r√©sultats"
  - 2-3 technologies en **gras** par comp√©tence (PAS PLUS)
  - Descriptions CONCISES, CLAIRES et PROFESSIONNELLES
  - Privil√©gier CLART√â et CONCISION sur la longueur
  
  "experiences_enrichies": [
    {{
      "periode": "2020-2023",
      "entreprise": "Nom entreprise",
      "poste": "Titre reformul√© selon JD",
      "responsabilites": [
        "Configuration **Open edX** incluant structuration et int√©gration avec **SharePoint** pour gestion contenus",
        "Automatisation processus documentaires via **Power Automate** et **Teams** pour am√©liorer efficacit√©"
      ],
      "environment": "**Open edX**, **SharePoint**, **Microsoft 365**, Teams, Power Automate, OneDrive, SQL"
    }}
  ],
  
  FORMAT OBLIGATOIRE (COPIER format comp√©tences):
  - Responsabilit√©s: Technologies **isol√©es** dans texte normal (ex: "Configuration **Tech1** incluant **Tech2** pour r√©sultats")
  - Environnement: Liste virgules avec 3-5 technologies **critiques** en gras, autres sans
  - JAMAIS phrases enti√®res en gras
  - Maximum 2-3 mots entre **ast√©risques**
  
  "score_matching": 45,
  "points_forts": ["ALWAYS in English: key strength 1", "ALWAYS in English: key strength 2"]
}}

üåç CRITICAL LANGUAGE REQUIREMENT:
- 'domaines_analyses' (domain names AND comments) MUST ALWAYS be in ENGLISH
- 'synthese_matching' MUST ALWAYS be in ENGLISH  
- 'points_forts' MUST ALWAYS be in ENGLISH
- Example domain: "SQL Data Extraction and Manipulation" NOT "Extraction de donn√©es SQL"
- Example comment: "‚ùå No demonstrated experience in SQL data extraction..." NOT "‚ùå Aucune exp√©rience..."
- Example synthesis: "Java senior profile unsuitable for .NET position..." NOT "Profil Java senior inadapt√©..."

CRITICAL SCORING RULES:
- 'domaines_analyses' MUST be completed with 5-8 domains totaling EXACTLY 100%
- BE STRICT on scoring - don't give points if candidate lacks the skill
- If stack mismatch (Java vs .NET), give 0 points, not 40-50

üî¥üî¥üî¥ V√âRIFICATION FINALE AVANT R√âPONSE (NON-N√âGOCIABLE) üî¥üî¥üî¥

AVANT de g√©n√©rer ta r√©ponse JSON, tu DOIS:

1Ô∏è‚É£ CALCULER LA SOMME:
   Somme = domaine1.score + domaine2.score + domaine3.score + ... + domaineN.score
   
2Ô∏è‚É£ V√âRIFIER:
   Si Somme ‚â† score_matching ‚Üí ERREUR ‚Üí RECALCULE TOUT
   
3Ô∏è‚É£ V√âRIFIER LES POIDS:
   Somme_Poids = domaine1.poids + domaine2.poids + ... + domaineN.poids
   Si Somme_Poids ‚â† 100 ‚Üí ERREUR ‚Üí RECALCULE TOUT
   
4Ô∏è‚É£ DOUBLE-CHECK:
   Pour chaque domaine: v√©rifie que score ‚â§ score_max
   Pour chaque domaine: v√©rifie que score_max = poids

EXEMPLE DE V√âRIFICATION:
Domaine 1: Stack .NET (40%) ‚Üí 0/40 points
Domaine 2: Cloud Azure (20%) ‚Üí 8/20 points  
Domaine 3: SQL Server (15%) ‚Üí 10/15 points
Domaine 4: DevOps (15%) ‚Üí 5/15 points
Domaine 5: Agile (10%) ‚Üí 7/10 points

V√©rification poids: 40+20+15+15+10 = 100 ‚úÖ
V√©rification score: 0+8+10+5+7 = 30 ‚úÖ
Donc: score_matching = 30 ‚úÖ

Si tu trouves une incoh√©rence ‚Üí RECALCULE TOUT depuis le d√©but

---

JOB DESCRIPTION:
{jd_text}

---

CV ACTUEL:
{cv_text}

---

IMPORTANT: JSON strict uniquement, sans commentaire ni balise."""

            response = client.messages.create(
                model="claude-sonnet-4-5-20250929",
                max_tokens=8000,
                messages=[{"role": "user", "content": prompt}]
            )
            
        except Exception as e:
            print(f">>> ERROR calling anthropic for enrichment: {repr(e)}", flush=True)
            return {}
        
        response_text = response.content[0].text.strip()
        
        # Nettoyer JSON
        if response_text.startswith('```json'):
            response_text = response_text[7:]
        if response_text.startswith('```'):
            response_text = response_text[3:]
        if response_text.endswith('```'):
            response_text = response_text[:-3]
        response_text = response_text.strip()
        
        try:
            enriched = json.loads(response_text)
            print(f"‚úÖ Enrichissement r√©ussi!")
            print(f"   Score matching: {enriched.get('score_matching', 0)}/100")
            print(f"   Domaines analys√©s: {len(enriched.get('domaines_analyses', []))}")
            print(f"   Mots-cl√©s en gras: {len(enriched.get('mots_cles_a_mettre_en_gras', []))}")
            
            if enriched.get('domaines_analyses'):
                print(f"\n   üìä D√©tail scoring:")
                for domaine in enriched['domaines_analyses']:
                    emoji = domaine.get('match', '')
                    if emoji == 'incompatible':
                        emoji = '‚ùå'
                    elif emoji == 'partiel':
                        emoji = '‚ö†Ô∏è'
                    elif emoji in ['bon', 'excellent']:
                        emoji = '‚úÖ'
                    print(f"      {emoji} {domaine.get('domaine', 'N/A')}: {domaine.get('score', 0)}/{domaine.get('score_max', 0)} ({domaine.get('poids', 0)}%)")
            
            # DEBUG: Afficher une responsabilit√© pour voir le format
            if enriched.get('experiences_enrichies'):
                first_exp = enriched['experiences_enrichies'][0]
                if first_exp.get('responsabilites'):
                    print(f"\nüîç DEBUG - Premi√®re responsabilit√© :")
                    print(f"   {first_exp['responsabilites'][0]}")
                if first_exp.get('environment'):
                    print(f"\nüîç DEBUG - Environnement :")
                    print(f"   {first_exp['environment']}")
            
            return enriched
        except json.JSONDecodeError as e:
            print(f"‚ö†Ô∏è Erreur JSON: {e}")
            return {}

    # ========================================
    # MODULE 4 : MAPPING TMC + RICHTEXT
    # ========================================
    
    def mdbold_to_richtext(self, s: str) -> RichText:
        """Convertit les **bold** markdown en RichText propre sans cascade de gras."""
        import re
        rt = RichText()
        pattern = re.compile(r'\*\*(.*?)\*\*')
        last_end = 0

        # Ajouter le texte avant chaque bloc en gras
        for match in pattern.finditer(s):
            if match.start() > last_end:
                rt.add(s[last_end:match.start()], bold=False, font='Arial')
            # Le texte entre **...** est en gras
            rt.add(match.group(1), bold=True, font='Arial')
            last_end = match.end()

        # Ajouter le texte apr√®s le dernier bloc
        if last_end < len(s):
            rt.add(s[last_end:], bold=False, font='Arial')

        return rt

    def map_to_tmc_structure(self, parsed_cv: Dict[str, Any], enriched_cv: Dict[str, Any], template_lang: str = 'FR') -> Dict[str, Any]:
        """Mapper les donn√©es enrichies vers la structure TMC"""
        print("üó∫Ô∏è  Mapping vers structure TMC...")
        
        # 1. PROFIL - Convertir en RichText pour supporter le gras (pas d'√©chappement)
        profil_brut = enriched_cv.get('profil_enrichi', parsed_cv.get('profil_resume', ''))
        profil = self.mdbold_to_richtext(profil_brut) if profil_brut else ''
        
        # 2. COMP√âTENCES - FORMAT CAT√âGORIS√â D√âTAILL√â
        competences_enrichies = enriched_cv.get('competences_enrichies', {})
        
        # Si competences_enrichies est un dict (nouveau format), l'utiliser directement
        if isinstance(competences_enrichies, dict):
            # Supprimer la cl√© "NOTE" si pr√©sente
            skills_categorized = {k: v for k, v in competences_enrichies.items() if k != 'NOTE' and isinstance(v, list)}
        else:
            # Fallback ancien format (liste simple)
            competences = competences_enrichies if isinstance(competences_enrichies, list) else parsed_cv.get('competences', [])
            skills_categorized = {
                'Comp√©tences techniques': competences[:8] if len(competences) >= 8 else competences,
                'Comp√©tences transversales': competences[8:12] if len(competences) > 8 else []
            }
            # Supprimer les cat√©gories vides
            skills_categorized = {k: v for k, v in skills_categorized.items() if v}
        
        # üî• Transformation en RichText pour le formatage (pas d'√©chappement)
        skills_categorized_doc = []
        for cat, skills in skills_categorized.items():
            rt_cat = RichText()
            rt_cat.add(cat, bold=True)
            rt_skills = [self.mdbold_to_richtext(s) for s in skills]
            skills_categorized_doc.append((rt_cat, rt_skills))
        
        # 3. EXP√âRIENCES - Texte simple pour les responsabilit√©s, RichText pour environnement
        experiences_enrichies = enriched_cv.get('experiences_enrichies', parsed_cv.get('experiences', []))
        work_experience = []
        
        for exp in experiences_enrichies:
            # GARDER les responsabilit√©s en TEXTE SIMPLE (pas RichText) - pas d'√©chappement
            responsabilites_text = [r for r in exp.get('responsabilites', [])]
            
            # Convertir l'environnement en RichText pour le gras - pas d'√©chappement
            environment_brut = exp.get('environment', '')
            environment_rt = self.mdbold_to_richtext(environment_brut) if environment_brut else ''
            
            work_exp = {
                'period': exp.get('periode', ''),
                'company': exp.get('entreprise', ''),
                'position': exp.get('poste', ''),
                'general_responsibilities': responsabilites_text,  # Texte simple
                'environment': environment_rt
            }
            work_experience.append(work_exp)
        
        # 4. FORMATION (avec d√©tails complets)
        formation = parsed_cv.get('formation', [])
        education = []
        for form in formation:
            education.append({
                'institution': form.get('institution', ''),
                'degree': form.get('diplome', ''),
                'graduation_year': form.get('annee', 'Date inconnue'),
                'country': form.get('pays', 'Canada'),
                'level': '',
                'title': form.get('diplome', '')
            })
        
        # 5. CERTIFICATIONS (avec mapping vers format template)
        certifications_raw = parsed_cv.get('certifications', [])
        certifications = []
        for cert in certifications_raw:
            certifications.append({
                'name': cert.get('nom', cert.get('name', '')),
                'institution': cert.get('organisme', cert.get('institution', '')),
                'year': str(cert.get('annee', cert.get('year', ''))),
                'country': cert.get('pays', cert.get('country', ''))
            })
        
        # 6. PROJETS
        projects = parsed_cv.get('projets', [])
        
        # 7. INFORMATIONS PERSONNELLES
        nom_complet = parsed_cv.get('nom_complet', '')
        
        # S√©parer pr√©nom et nom
        parts = nom_complet.split() if nom_complet else []
        if len(parts) >= 2:
            first_name = parts[0]
            last_name = ' '.join(parts[1:])
        elif len(parts) == 1:
            first_name = parts[0]
            last_name = ''
        else:
            first_name = 'Pr√©nom'
            last_name = 'Nom'
        
        titre_professionnel = enriched_cv.get('titre_professionnel_enrichi', parsed_cv.get('titre_professionnel', ''))
        lieu_residence = parsed_cv.get('lieu_residence', 'Montr√©al, Canada')
        langues_list = parsed_cv.get('langues', ['Fran√ßais', 'Anglais'])
        
        # Traduire les langues selon le template
        if template_lang == 'FR':
            # Si template FR, traduire de l'anglais vers le fran√ßais
            langue_map = {
                'English': 'Anglais',
                'French': 'Fran√ßais',
                'Hebrew': 'H√©breu',
                'Russian': 'Russe',
                'Spanish': 'Espagnol',
                'German': 'Allemand',
                'Italian': 'Italien',
                'Portuguese': 'Portugais',
                'Chinese': 'Chinois',
                'Japanese': 'Japonais',
                'Arabic': 'Arabe'
            }
            langues_list = [langue_map.get(lang, lang) for lang in langues_list]
        
        langues = ', '.join(langues_list)
        
        context = {
            # Pour le header (minuscules) - PAS d'√©chappement
            'first_name': first_name,
            'last_name': last_name,
            'title': titre_professionnel,
            
            # Pour la page 1 (MAJUSCULES) - PAS d'√©chappement
            'FIRST_NAME': first_name.upper(),
            'LAST_NAME': last_name.upper(),
            'TITLE': titre_professionnel,
            'RESIDENCY': lieu_residence,
            'LANGUAGES': langues,
            
            # AUSSI en minuscules pour compatibilit√© template
            'residency': lieu_residence,
            'languages': langues,
            
            # Reste du CV
            'summary': profil,
            'skills_categorized': skills_categorized,
            'skills_categorized_doc': skills_categorized_doc,  # üî• Version RichText pour le template
            'work_experience': work_experience,
            'education': education,
            'projects': projects,
            'certifications': certifications
        }
        
        print(f"‚úÖ Mapping termin√©!")
        print(f"   Nom: [ANONYMIZED]")
        print(f"   Titre: {titre_professionnel}")
        print(f"   Langues: {langues}")
        print(f"   Profil: RichText g√©n√©r√©")
        total_competences = sum(len(v) for v in skills_categorized.values() if isinstance(v, list))
        print(f"   Cat√©gories: {len(skills_categorized)}")
        print(f"   Comp√©tences: {total_competences}")
        print(f"   Exp√©riences: {len(work_experience)}")
        
        return context

    # ========================================
    # MODULE 5 : G√âN√âRATION DOCX TMC
    # ========================================
    
    def generate_tmc_docx(self, context: Dict[str, Any], output_path: str, template_path: str = "TMC_NA_template_FR.docx"):
        """G√©n√©rer le CV TMC final avec docxtpl"""
        print(f"üìù G√©n√©ration du CV TMC: {output_path}")
        print(f"   Template utilis√©: {template_path}")
        
        # Cr√©er environnement Jinja2 avec filtre pairwise
        jinja_env = jinja2.Environment()
        
        def pairwise(iterable):
            items = list(iterable)
            result = []
            for i in range(0, len(items), 2):
                if i + 1 < len(items):
                    result.append((items[i], items[i + 1]))
                else:
                    result.append((items[i], ''))
            return result
        
        jinja_env.filters['pairwise'] = pairwise
        
        # üî• Ajouter la fonction r pour RichText dans le contexte
        context['r'] = lambda x: x
        
        # Charger le template TMC
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"‚ùå Template TMC introuvable: {template_path}")
        
        doc = DocxTemplate(template_path)
        
        # Rendre le document
        doc.render(context, jinja_env)
        
        # Sauvegarder
        doc.save(output_path)
        print(f"‚úÖ CV TMC g√©n√©r√© avec succ√®s!")

    def apply_bold_post_processing(self, docx_path: str, keywords: list):
        """Post-traiter le document pour mettre en gras les technologies dans les tableaux"""
        print(f"üé® Application du gras sur les technologies...")
        
        from docx import Document as DocxDocument
        from docx.shared import RGBColor
        import re
        
        doc = DocxDocument(docx_path)
        modifications = 0
        
        print(f"   Recherche des **mot** dans le document...")
        
        def apply_bold_to_runs(paragraph):
            """Trouve **mot** et met en gras UNIQUEMENT ce mot"""
            text = paragraph.text
            if '**' not in text:
                return 0
            
            changes = 0
            # Pattern pour trouver **mot**
            pattern = re.compile(r'\*\*([^*]+)\*\*')
            
            # Reconstituer le paragraphe avec le bon formatage
            matches = list(pattern.finditer(text))
            if not matches:
                return 0
            
            # Supprimer tous les runs existants
            for run in paragraph.runs:
                run._element.getparent().remove(run._element)
            
            # Reconstruire avec le bon formatage
            last_end = 0
            for match in matches:
                # Texte normal avant
                if match.start() > last_end:
                    run = paragraph.add_run(text[last_end:match.start()])
                    run.bold = False
                    run.font.name = 'Arial'
                
                # Texte en gras
                run = paragraph.add_run(match.group(1))
                run.bold = True
                run.font.name = 'Arial'
                changes += 1
                
                last_end = match.end()
            
            # Texte normal apr√®s
            if last_end < len(text):
                run = paragraph.add_run(text[last_end:])
                run.bold = False
                run.font.name = 'Arial'
            
            return changes
        
        # Parcourir TOUS les tableaux (o√π sont les exp√©riences)
        print("   üìã Traitement des tableaux...")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        modifications += apply_bold_to_runs(paragraph)
        
        # Parcourir aussi les paragraphes normaux
        print("   üìù Traitement des paragraphes...")
        for paragraph in doc.paragraphs:
            modifications += apply_bold_to_runs(paragraph)
        
        # Sauvegarder
        doc.save(docx_path)
        if modifications > 0:
            print(f"‚úÖ {modifications} mots mis en gras")
        else:
            print(f"‚ö†Ô∏è Aucun **mot** trouv√©")
        
        return modifications
        
def main():
    """Point d'entr√©e CLI"""
    import argparse
    
    parser = argparse.ArgumentParser(description='TMC Universal CV Enricher')
    parser.add_argument('cv_path', help='Chemin du CV (PDF, Word, etc.)')
    parser.add_argument('jd_path', help='Chemin de la Job Description')
    parser.add_argument('--output', '-o', default='cv_enriched_tmc.docx', help='Fichier de sortie')
    
    args = parser.parse_args()
    
    try:
        enricher = TMCUniversalEnricher()
        
        print("\nüöÄ TMC UNIVERSAL CV ENRICHER")
        print("=" * 60)
        
        # MODULE 1: Extraction
        print("\n[1/5] Extraction du CV...")
        cv_text = enricher.extract_cv_text(args.cv_path)
        print(f"      ‚úÖ {len(cv_text)} caract√®res extraits")
        
        # MODULE 2: Parsing
        print("\n[2/5] Parsing intelligent...")
        parsed_cv = enricher.parse_cv_with_claude(cv_text)
        
        # MODULE 3: Enrichissement
        print("\n[3/5] Enrichissement avec IA...")
        jd_text = enricher.read_job_description(args.jd_path)
        enriched_cv = enricher.enrich_cv_with_prompt(parsed_cv, jd_text)
        
        # MODULE 4: Mapping TMC
        print("\n[4/5] Mapping structure TMC...")
        tmc_context = enricher.map_to_tmc_structure(parsed_cv, enriched_cv)
        
        # MODULE 5: G√©n√©ration
        print("\n[5/5] G√©n√©ration CV final...")
        enricher.generate_tmc_docx(tmc_context, args.output)
        
        # POST-PROCESSING: Application du gras
        print("\n[POST] Application du gras sur mots-cl√©s...")
        keywords = enriched_cv.get('mots_cles_a_mettre_en_gras', [])
        print(f"   Mots-cl√©s √† mettre en gras: {keywords}")
        
        if keywords:
            result = enricher.apply_bold_post_processing(args.output, keywords)
            if result == 0:
                print("   ‚ö†Ô∏è AUCUN mot-cl√© n'a √©t√© mis en gras!")
                print("   V√©rifiez que les mots-cl√©s sont bien dans le CV")
        else:
            print("   ‚ö†Ô∏è Aucun mot-cl√© retourn√© par l'IA")
        
        # R√âSUM√â FINAL
        print("\n" + "=" * 60)
        print("üéâ ENRICHISSEMENT TERMIN√â!")
        print("=" * 60)
        print(f"üìä Score matching: {enriched_cv.get('score_matching', 0)}/100")
        
        # Afficher les domaines analys√©s
        if enriched_cv.get('domaines_analyses'):
            print(f"\nüìä Analyse par domaine:")
            for domaine in enriched_cv['domaines_analyses']:
                match = domaine.get('match', '')
                emoji = '‚ùå' if match == 'incompatible' else '‚ö†Ô∏è' if match == 'partiel' else '‚úÖ'
                print(f"   {emoji} {domaine.get('domaine', 'N/A')}: {domaine.get('score', 0)}/{domaine.get('score_max', 0)} pts ({domaine.get('poids', 0)}%)")
                print(f"      ‚Üí {domaine.get('commentaire', 'N/A')}")
        
        if enriched_cv.get('synthese_matching'):
            print(f"\nüí¨ Synth√®se: {enriched_cv['synthese_matching']}")
        
        print(f"\nüí™ Points forts:")
        for pf in enriched_cv.get('points_forts', [])[:3]:
            print(f"   ‚Ä¢ {pf}")
        print(f"\nüìÑ Fichier g√©n√©r√©: {args.output}")
        print("=" * 60)
        
    except Exception as e:
        print(f"\n‚ùå ERREUR: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
